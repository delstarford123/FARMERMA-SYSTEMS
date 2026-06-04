from gevent import monkey
monkey.patch_all()

import os
from flask import Flask
from flask_socketio import SocketIO
# ... rest of your imports
import os
import json
import base64
import stripe
import requests
import threading
import smtplib
import time
import uuid # <-- Essential for generating unique cloud filenames
from datetime import datetime, timedelta, timezone
from functools import wraps 
from email.message import EmailMessage
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from werkzeug.exceptions import NotFound
from logic import analyze_weather_and_generate_alerts, update_firebase_alerts
# Flask & SocketIO
from flask import Flask, render_template, request, redirect, url_for, jsonify, session, flash, send_from_directory, abort
from flask_mail import Mail, Message
from flask_socketio import SocketIO, emit, join_room, leave_room

# Twilio & Google GenAI Imports
from twilio.twiml.messaging_response import MessagingResponse
from google import genai
from google.genai import types
from docx import Document
from io import BytesIO

# Firebase Imports
import firebase_admin
from firebase_admin import credentials, auth, db as firebase_db, storage # <-- Essential: Added 'storage'

# Internal Project Imports
from models import db as sqlalchemy_db, User, MarketData, Transaction 
from ai_logic.ai_engine import generate_price_forecast
from mpesa import initiate_stk_push
from pesapal_helper import get_pesapal_token, register_pesapal_ipn, submit_pesapal_order, get_pesapal_transaction_status

# APScheduler Setup
from apscheduler.schedulers.background import BackgroundScheduler

# ==========================================
# 1. INITIALIZATION & APP CONFIGURATION
# ==========================================
load_dotenv()

app = Flask(__name__)
app.secret_key = 'delstarford_works_secret_key' 
app.config['PERMANENT_SESSION_LIFETIME'] = 86400 # 24 hours

# Flask-Mail Configuration
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = ('Farmerman Systems', os.environ.get('MAIL_USERNAME'))

mail = Mail(app)

# APScheduler Initialization
scheduler = BackgroundScheduler()
scheduler.start()

# ==========================================
# SUBSCRIPTION RENEWAL REMINDERS
# ==========================================
def check_subscription_expirations():
    """Daily task to send payment reminders 3 days before expiry."""
    with app.app_context():
        try:
            eat_tz = timezone(timedelta(hours=3))
            now = datetime.now(eat_tz)
            reminder_date = (now + timedelta(days=3)).strftime("%Y-%m-%d")
            
            all_users = rtdb.reference('users').get() or {}
            for uid, data in all_users.items():
                expiry_str = data.get('subscription_expiry')
                if expiry_str:
                    # Check if the date matches our 3-day window
                    if expiry_str.startswith(reminder_date):
                        email = data.get('email')
                        name = data.get('full_name', 'Valued Farmer')
                        plan = data.get('subscription_tier', 'Seed').capitalize()
                        
                        # Send Renewal Email
                        msg = Message(f"⚠️ Your ZIMBOT {plan} Package expires soon!", recipients=[email])
                        msg.body = f"Hello {name},\n\nYour {plan} subscription will expire in 3 days. Please visit your dashboard to renew and maintain uninterrupted access to market intelligence.\n\nRenew here: {request.host_url}pricing\n\nBest regards,\nZimBot Team"
                        mail.send(msg)
                        print(f"Renewal reminder sent to {email}")
        except Exception as e:
            print(f"Reminder Job Error: {e}")

# Add the job to run every 24 hours
scheduler.add_job(func=check_subscription_expirations, trigger='interval', hours=24)

# ==========================================
# WEEKLY HARVEST INTELLIGENCE REPORTS
# ==========================================
def generate_harvest_report(user_name):
    """Generates a professional Word document report for Harvest subscribers."""
    doc = Document()
    doc.add_heading('Harvest Intelligence: Weekly Strategic Report', 0)
    
    doc.add_paragraph(f"Prepared for: {user_name}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    
    doc.add_heading('1. Market Overview (Zimbabwe)', level=1)
    
    # Fetch live market data for the report
    market_ref = rtdb.reference('market_data')
    all_data = market_ref.get() or {}
    zim_items = [item for item in all_data.values() if item.get('country') == 'Zimbabwe']
    
    if zim_items:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Commodity'
        hdr_cells[1].text = 'Price'
        hdr_cells[2].text = 'Trend'
        hdr_cells[3].text = 'Market'
        
        for item in zim_items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('commodity'))
            row_cells[1].text = f"{item.get('currency')} {item.get('price')} / {item.get('unit')}"
            row_cells[2].text = "Rising ▲" if item.get('trend') == 'up' else "Falling ▼" if item.get('trend') == 'down' else "Stable"
            row_cells[3].text = str(item.get('region'))
    else:
        doc.add_paragraph("No live market data available for this period.")

    doc.add_heading('2. Strategic Action Plan', level=1)
    doc.add_paragraph(
        "Based on current market volatility, Harvest subscribers are advised to prioritize storage for grains showing a downward trend. "
        "For commodities in a rising trend (▲), incremental selling is recommended to capture the peak ROI while maintaining supply continuity."
    )
    
    doc.add_heading('3. Intelligence Forecast', level=1)
    doc.add_paragraph(
        "Our data models predict a tightening of supply in the coming 14 days due to regional logistics shifts. "
        "Buyers should secure forward contracts now to lock in current rates."
    )
    
    doc.add_paragraph("\n\nThank you for choosing ZimBot Harvest Intelligence.")
    doc.add_paragraph("ZimBot Team - Your Professional Agricultural Partner.")
    
    target = BytesIO()
    doc.save(target)
    target.seek(0)
    return target

def send_friday_harvest_reports():
    """Weekly task to send Word reports to Harvest subscribers every Friday at 6 PM."""
    with app.app_context():
        try:
            print("🚀 Starting Weekly Harvest Report distribution...")
            all_users = rtdb.reference('users').get() or {}
            
            # Filter for active Harvest subscribers (Package 3)
            harvest_users = [
                {'email': u.get('email'), 'name': u.get('full_name', 'Valued Farmer')}
                for u in all_users.values()
                if u.get('subscription_tier') == 'harvest' and u.get('subscription_status') == 'active'
            ]
            
            for user in harvest_users:
                email = user['email']
                name = user['name']
                
                if not email: continue
                
                # Generate Report
                report_buffer = generate_harvest_report(name)
                
                # Create Email
                msg = Message(
                    subject=f"📊 Your Weekly Harvest Intelligence Report - {datetime.now().strftime('%b %d')}",
                    recipients=[email]
                )
                msg.body = f"Hello {name},\n\nPlease find attached your exclusive Weekly Harvest Intelligence Report. This strategic document contains live market trends and action plans tailored for your agribusiness success.\n\nBest regards,\nZimBot Team"
                
                # Attach Word Doc
                msg.attach(
                    filename=f"Harvest_Report_{datetime.now().strftime('%Y_%m_%d')}.docx",
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    data=report_buffer.read()
                )
                
                mail.send(msg)
                print(f"✅ Harvest Report sent to {email}")
                
        except Exception as e:
            print(f"❌ Harvest Report distribution failed: {e}")

# Schedule the Harvest Report for every Friday at 18:00 (6 PM)
scheduler.add_job(
    func=send_friday_harvest_reports, 
    trigger='cron', 
    day_of_week='fri', 
    hour=18, 
    minute=0
)

# Database Config
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///farmerman.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
sqlalchemy_db.init_app(app)

# Initialize SocketIO for real-time chat
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='gevent')

# Dictionary to track who is currently online { 'user_id': 'socket_id' }
online_users = {}

# Folder for legacy chat media (Fallback if cloud fails)
CHAT_MEDIA_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'chat_media')
os.makedirs(CHAT_MEDIA_FOLDER, exist_ok=True)

# ==========================================
# CONSTANTS & UPLOAD CONFIGURATIONS
# ==========================================
FIREBASE_WEB_API_KEY = "AIzaSyDy41jUJ8h7zYE9Ocj7pPNGGXCq5RRbN-s"
ALLOWED_TRAINING_EXTENSIONS = {'mp4', 'pdf', 'png', 'jpg', 'jpeg', 'docx', 'mp3', 'wav', 'avi', 'webp'}
PREMIUM_CONTENT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'premium_content')
os.makedirs(PREMIUM_CONTENT_FOLDER, exist_ok=True) 

def allowed_training_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_TRAINING_EXTENSIONS

# ==========================================
# FIREBASE SECURE CONNECTION
# ==========================================
rtdb = None
try:
    render_path = '/etc/secrets/serviceAccountKey.json'
    local_path = 'serviceAccountKey.json'
    cert_path = render_path if os.path.exists(render_path) else local_path

    if not firebase_admin._apps:
        cred = credentials.Certificate(cert_path)
        firebase_admin.initialize_app(cred, {
            'databaseURL': 'https://farmerman-systems-default-rtdb.firebaseio.com/',
            'storageBucket': 'farmerman-systems.firebasestorage.app'
        })
    
    rtdb = firebase_db 
    print(f"Firebase securely initialized using: {cert_path}")

    # CONNECTION TEST
    test_fetch = rtdb.reference('users').get()
    if test_fetch:
        print(f"Connection Verified: Found {len(test_fetch)} records in 'users' node.")
    else:
        print("Warning: Connection successful but 'users' node appears empty.")

except Exception as e:
    print(f"Firebase Initialization Error: {e}")

with app.app_context():
    sqlalchemy_db.create_all()


# ==========================================
# FIREBASE CLOUD STORAGE UPLOADER (NEW!)
# ==========================================
def upload_to_firebase_storage(file_obj, folder_name):
    """Uploads a file to Firebase Cloud Storage and returns an unguessable public URL."""
    try:
        bucket = storage.bucket()
        
        # 1. Generate a secure, unique filename (e.g., 123e4567-e89b.png)
        extension = file_obj.filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4()}.{extension}"
        blob_path = f"{folder_name}/{unique_filename}"
        
        # 2. Upload to Firebase
        blob = bucket.blob(blob_path)
        file_obj.seek(0) # Ensure we read from the beginning of the file
        blob.upload_from_file(file_obj, content_type=file_obj.content_type)
        
        # 3. Make the file readable by the web browser
        blob.make_public()
        
        # 4. Return the permanent cloud URL
        return blob.public_url
        
    except Exception as e:
        print(f"Firebase Storage Error: {e}")
        return None


# ==========================================
# ASYNC EMAIL FUNCTIONS
# ==========================================
def send_async_emails(user_email, admin_email, user_msg_html, admin_msg_html, name, message_body, inquiry_subject):
    """Background worker for Contact Forms using smtplib."""
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server: 
            server.login(app.config['MAIL_USERNAME'], app.config['MAIL_PASSWORD'])
            
            user_msg = EmailMessage()
            user_msg['Subject'] = "We received your message - Farmerman Systems"
            user_msg['From'] = f"Farmerman Support <{app.config['MAIL_USERNAME']}>"
            user_msg['To'] = user_email
            user_msg.set_content("Thank you for contacting Farmerman Systems. We will get back to you shortly.")
            user_msg.add_alternative(user_msg_html, subtype='html')
            server.send_message(user_msg)
            
            admin_msg = EmailMessage()
            admin_msg['Subject'] = f"🚨 {inquiry_subject} Inquiry from {name}"
            admin_msg['From'] = f"Farmerman Server <{app.config['MAIL_USERNAME']}>"
            admin_msg['To'] = admin_email 
            admin_msg.set_content(f"New {inquiry_subject} message from {name}: {message_body}")
            admin_msg.add_alternative(admin_msg_html, subtype='html')
            server.send_message(admin_msg)
            
    except Exception as e:
        print(f"Failed to send background emails: {e}")

def send_drip_followup(user_email, name):
    """Background worker for the 3-Day Drip Campaign using Flask-Mail."""
    with app.app_context(): 
        try:
            user_data = rtdb.reference('users').order_by_child('email').equal_to(user_email).get()
            if user_data:
                uid = list(user_data.keys())[0]
                current_tier = user_data[uid].get('subscription_tier', 'free')

                if current_tier == 'free':
                    msg = Message("Still guessing market prices? 📈", recipients=[user_email])
                    msg.html = render_template('emails/drip_upgrade.html', name=name)
                    mail.send(msg)
                    print(f"Drip email successfully sent to {user_email}")
        except Exception as e:
            print(f"Failed to send Drip Campaign email: {e}")

# ==========================================
# CONTEXT PROCESSORS
# ==========================================
@app.context_processor
def inject_site_content():
    def get_site_content(page_id):
        try:
            return rtdb.reference(f'site_content/{page_id}').get()
        except Exception:
            return None
    return dict(get_site_content=get_site_content)

@app.context_processor
def inject_global_vars():
    """Injects secure .env variables globally into all Jinja templates (e.g., base.html)."""
    return dict(
        weather_api_key=os.environ.get('WEATHER_API_KEY')
    )

# ==========================================
# ENTERPRISE SECURITY DECORATORS
# ==========================================
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash("Please log in to access this page.", "warning")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session or session.get('role') != 'admin':
            flash("Access Denied: Administrator privileges required.", "danger")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def tutor_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        user_role = session.get('role', session.get('user_role'))
        if user_role not in ['tutor', 'admin']:
            flash("Access denied: This area is reserved for Tutors.", "danger")
            return redirect(url_for('insights')) # <-- New standard landing page
        return f(*args, **kwargs)
    return decorated_function

def premium_required(f):
    """The Gatekeeper for Premium Intelligence and Pro Academy content."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash("Please login to access Premium content.", "warning")
            return redirect(url_for('login'))
        
        user_role = session.get('role', session.get('user_role'))
        user_tier = session.get('tier', session.get('subscription_tier', 'free'))
        
        # Admins and Tutors bypass billing
        if user_role in ['admin', 'tutor']:
            return f(*args, **kwargs)

        # Only Paid tiers get through
        if user_tier not in ['premium', 'pro', 'enterprise', 'bronze', 'silver', 'gold']:
            flash("Upgrade Required: This is Premium Content.", "warning")
            return redirect(url_for('pricing'))
        
        return f(*args, **kwargs)
    return decorated_function

def token_required(f):
    """API token verification for external endpoints."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token: return jsonify({"error": "Token missing"}), 401
        try:
            if token.startswith("Bearer "): token = token.split(" ")[1]
            request.user = auth.verify_id_token(token)
        except Exception: return jsonify({"error": "Invalid token"}), 401
        return f(*args, **kwargs)
    return decorated_function

def token_admin_required(f):
    """Admin-only API token verification."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token: return jsonify({"error": "Token missing"}), 401
        try:
            if token.startswith("Bearer "): token = token.split(" ")[1]
            decoded_token = auth.verify_id_token(token)
            uid = decoded_token['uid']
            user_data = rtdb.reference(f'users/{uid}').get()
            if not user_data or user_data.get('role') != 'admin':
                return jsonify({"error": "Admin required"}), 403
            request.uid = uid
        except Exception: return jsonify({"error": "Invalid token"}), 401
        return f(*args, **kwargs)
    return decorated_function


# ==========================================
# BACKGROUND EMAIL WORKER
# ==========================================
def send_welcome_email(user_email, name, role):
    """Sends the welcome email in the background so the user doesn't wait."""
    with app.app_context():
        try:
            # Personalize subject based on the new roles
            if role == 'tutor':
                subject = "Welcome to the Faculty!"
            elif role == 'seller':
                subject = "Ready to scale your agribusiness? 🚀"
            else: # buyer or legacy client
                subject = "Your Market Intelligence is Ready!"
                
            msg = Message(subject, recipients=[user_email])
            
            template = 'emails/welcome_tutor.html' if role == 'tutor' else 'emails/welcome_client.html'
            msg.html = render_template(template, name=name)
            
            mail.send(msg)
            print(f"Welcome email successfully sent to {user_email}")
        except Exception as e:
            print(f"Failed to send welcome email: {e}")



# ==========================================
# ROBUST AUTHENTICATION ROUTES
# ==========================================
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form.get('email').strip()
        password = request.form.get('password')
        full_name = request.form.get('fullName').strip()
        organization = request.form.get('organization', '').strip()
        
        # 1. SECURITY: Capture and sanitize the role (UPDATED FOR BUYER/SELLER/TUTOR)
        selected_role = request.form.get('role', 'buyer').strip().lower()
        if selected_role not in ['buyer', 'seller', 'tutor']:
            selected_role = 'buyer'

        try:
            # 2. Create Auth User
            user = auth.create_user(email=email, password=password, display_name=full_name)
            
            # 3. Initialize profile in RTDB
            rtdb.reference(f'users/{user.uid}').set({
                'uid': user.uid,
                'full_name': full_name, 
                'email': email, 
                'organization': organization,
                'role': selected_role, 
                'subscription_tier': 'free', 
                'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            
            # 4. Fire off the Welcome Email
            threading.Thread(target=send_welcome_email, args=(email, full_name, selected_role)).start()
            
            # 5. Schedule the Drip Campaign
            run_time = datetime.now() + timedelta(days=3)
            scheduler.add_job(
                func=send_drip_followup,
                trigger='date',
                run_date=run_time,
                args=[email, full_name]
            )

            # 6. Redirect to Login
            flash("Account created successfully! Check your inbox for the welcome email.", "success")
            return redirect(url_for('login'))
            
        except Exception as e: 
            flash(f"Registration Error: {str(e)}", "danger")
            
    return render_template('register.html')
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email').strip()
        password = request.form.get('password')
        
        try:
            # 1. Firebase Auth REST API
            request_ref = f"https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key={FIREBASE_WEB_API_KEY}"
            data = {"email": email, "password": password, "returnSecureToken": True}
            req = requests.post(request_ref, json=data)
            req.raise_for_status() 
            
            user = req.json()
            uid = user['localId'] 
            
            # 2. Database Fetch & Auto-Heal
            user_ref = rtdb.reference(f'users/{uid}')
            user_data = user_ref.get()
            
            if user_data:
                # Fallback to 'buyer' if role is empty to match your new frontend
                raw_role = str(user_data.get('role', 'buyer')).strip().lower()
                raw_tier = str(user_data.get('subscription_tier', 'free')).strip().lower()
            else:
                raw_role = 'buyer'
                raw_tier = 'free'
                user_ref.set({'email': email, 'role': raw_role, 'subscription_tier': raw_tier, 'uid': uid})

            # 3. Secure the Session
            session.clear() 
            session.permanent = True 
            session['user_id'] = uid
            session['user_email'] = email
            session['role'] = raw_role
            session['tier'] = raw_tier 
            session['subscription_tier'] = raw_tier 

            # 4. Smart Redirects based on Role
            if raw_role == 'admin':
                flash("Welcome back, Administrator!", "success")
                return redirect(url_for('subscriber_management')) 
                
            elif raw_role == 'tutor':
                flash("Welcome to the Faculty Portal!", "success")
                # Usually tutors go to their dashboard, or keep it as market_intelligence if you prefer
                return redirect(url_for('tutor_dashboard')) 
                
            elif raw_role in ['buyer', 'seller', 'client']:
                # DIRECT BUYERS & SELLERS TO INSIGHTS & RESEARCH
                # (We keep 'client' in the list just in case you have older legacy users in the database)
                flash(f"Authentication successful. Welcome to your {raw_role.capitalize()} portal!", "success")
                return redirect(url_for('insights'))
                
            else:
                # Catch-all fallback
                flash("Authentication successful. Welcome to your portal!", "success")
                return redirect(url_for('insights'))
            
        except requests.exceptions.HTTPError:
            flash("Invalid email or password. Please try again.", "danger")
        except Exception as e:
            print(f"Login System Error: {e}")
            flash("System error during login. Check server console.", "danger")
            
    return render_template('login.html')



@app.route('/reset-password', methods=['GET', 'POST'])
def reset_password():
    if request.method == 'POST':
        email = request.form.get('email').strip()
        try:
            request_ref = f"https://identitytoolkit.googleapis.com/v1/accounts:sendOobCode?key={FIREBASE_WEB_API_KEY}"
            headers = {"content-type": "application/json; charset=UTF-8"}
            data = {"requestType": "PASSWORD_RESET", "email": email}
            req = requests.post(request_ref, headers=headers, json=data)
            req.raise_for_status() 
            flash(f"A password reset link has been sent to {email}. Please check your inbox.", "success")
            return redirect(url_for('login'))
        except requests.exceptions.HTTPError:
            error_data = req.json().get('error', {}).get('message', '')
            if error_data == "EMAIL_NOT_FOUND":
                flash("No account is registered with that email address.", "warning")
            else:
                flash("Failed to send reset email. Please try again.", "danger")
        except Exception as e:
            flash("System error. Please contact support.", "danger")
            
    return render_template('reset_password.html')

@app.route('/logout')
def logout():
    role = session.get('role')
    session.clear()
    
    if role == 'admin':
        flash("Admin session securely terminated.", "success")
    else:
        flash("You have been securely logged out.", "info")
        
    return redirect(url_for('login'))

# ==========================================
# CLIENT DASHBOARD & SETTINGS
# ==========================================
@app.route('/dashboard')
@login_required
def dashboard():
    uid = session['user_id']
    try:
        profile = rtdb.reference(f'users/{uid}').get() or {}
        if 'full_name' not in profile: profile['full_name'] = 'Valued Farmer'
        return render_template('dashboard.html', profile=profile)
    except Exception as e:
        return render_template('dashboard.html', profile={'full_name': 'User'})

@app.route('/billing')
@login_required
def billing_history():
    user_id = session.get('user_id')
    profile = rtdb.reference(f'users/{user_id}').get() or {}
    current_plan = profile.get('subscription_tier', session.get('subscription_tier', 'free'))
    txns_ref = rtdb.reference(f'completed_transactions/{user_id}').get()
    transactions_list = sorted(txns_ref.values(), key=lambda x: x.get('date', ''), reverse=True) if txns_ref else []
        
    return render_template(
        'payments/billing_history.html', 
        profile=profile, 
        transactions=transactions_list, 
        current_plan=current_plan
    )

@app.route('/payment-failed')
@login_required
def payment_failed():
    error_message = request.args.get('msg', None)
    return render_template('payments/payment_failed.html', error_message=error_message)

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def account_settings(): 
    uid = session['user_id']
    user_ref = rtdb.reference(f'users/{uid}')
    
    if request.method == 'POST':
        user_ref.update({
            'full_name': request.form.get('full_name'),
            'phone': request.form.get('phone'),
            'location': request.form.get('location')
        })
        flash("Settings updated successfully!", "success")
        return redirect(url_for('account_settings'))
        
    profile_data = user_ref.get() or {'full_name': 'User', 'email': session.get('user_email', '')}
    return render_template('accounts&subscription settings.html', profile=profile_data)
      
# ==========================================
# ADMIN HUB (Fully Protected)
# ==========================================
@app.route('/admin/dashboard')
@admin_required
def admin_dashboard(): 
    users = rtdb.reference('users').get() or {}
    market = rtdb.reference('market_data').get() or {}
    txns = rtdb.reference('completed_transactions').get() or {}
    
    # 1. Existing Revenue & Subscriptions Logic
    rev = 0.0; recent = []
    for uid, u_txns in txns.items():
        u_info = users.get(uid, {})
        for t in u_txns.values():
            rev += float(t.get('amount', 0))
            recent.append({'name': u_info.get('full_name', 'User'), 'date': t.get('date', ''), 'plan': t.get('plan', 'Pro'), 'email': u_info.get('email', 'N/A')})
            
    recent.sort(key=lambda x: x['date'], reverse=True)
    
    # 2. NEW: Global Banking System Analytics
    banking_groups = rtdb.reference('banking_groups').get() or {}
    total_groups = len(banking_groups)
    
    banking_accounts = rtdb.reference('banking_accounts').get() or {}
    total_deposits = 0.0
    for acc in banking_accounts.values():
        total_deposits += float(acc.get('emergency_fund', 0.0))
        for grp_balance in acc.get('standard_savings', {}).values():
            total_deposits += float(grp_balance)
            
    # 3. NEW: Pending Loan Requests
    all_loans = rtdb.reference('banking_loans').get() or {}
    pending_loans = []
    for uid, user_loans in all_loans.items():
        if isinstance(user_loans, dict):
            for loan_id, loan_data in user_loans.items():
                if loan_data.get('status') == 'Pending Review':
                    loan_data['loan_id'] = loan_id
                    loan_data['uid'] = uid
                    
                    # Attach user contact info so admin can call them
                    user_profile = users.get(uid, {})
                    loan_data['user_name'] = user_profile.get('full_name', 'Unknown')
                    loan_data['user_email'] = user_profile.get('email', 'No Email')
                    # Prefer the phone number typed in the loan request, fallback to profile phone if missing
                    loan_data['user_phone'] = loan_data.get('phone_number', user_profile.get('phone', 'No Phone Provided'))
                    
                    pending_loans.append(loan_data)
                    
    pending_loans.sort(key=lambda x: x.get('requested_at', ''), reverse=True)

    return render_template(
        'admin dashboard.html', 
        total_subscribers=len(users), 
        active_feeds=len(market), 
        total_revenue=rev, 
        recent_transactions=recent[:5],
        total_groups=total_groups,
        total_deposits=total_deposits,
        pending_loans=pending_loans
    )
    
@app.route('/admin/process-loan', methods=['POST'])
@admin_required
def admin_process_loan():
    uid = request.form.get('uid')
    loan_id = request.form.get('loan_id')
    action = request.form.get('action') # 'Approved' or 'Denied'
    
    try:
        rtdb.reference(f'banking_loans/{uid}/{loan_id}').update({
            'status': action,
            'processed_at': datetime.now(timezone(timedelta(hours=3))).strftime("%Y-%m-%d %H:%M:%S")
        })
        flash(f"Loan successfully {action}.", "success")
    except Exception as e:
        flash(f"Error processing loan: {e}", "danger")
        
    return redirect(url_for('admin_dashboard'))
   
@app.route('/admin/upload-training-media', methods=['GET', 'POST'])
@admin_required
def admin_upload_training():
    if request.method == 'POST':
        file = request.files.get('file')
        description = request.form.get('description', 'No description provided.')
        category = request.form.get('category', 'agripreneur') 
        
        if not file or file.filename == '':
            flash('No file selected.', 'warning')
            return redirect(request.url)
            
        if file and allowed_training_file(file.filename):
            filename = secure_filename(file.filename)
            save_path = os.path.join(PREMIUM_CONTENT_FOLDER, filename)
            file_extension = filename.rsplit('.', 1)[1].lower()
            
            try:
                file.save(save_path) 
                rtdb.reference('training_content').push({
                    'filename': filename, 'description': description, 'category': category,
                    'file_type': file_extension, 'upload_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
                flash(f'Success! {filename} published to academy.', 'success')
            except Exception as e:
                flash(f'Error saving file: {e}', 'danger')
            return redirect(url_for('admin_upload_training'))
        else:
            flash(f'Invalid file type.', 'danger')

    files_metadata = list((rtdb.reference('training_content').get() or {}).values())
    return render_template('admin_training_upload.html', files_metadata=files_metadata)
@app.route('/admin/data-manager', methods=['GET', 'POST'])
@admin_required
def market_data_manager():
    if request.method == 'POST':
        # Safely convert price to float for graphing
        try:
            numeric_price = float(request.form.get('price', 0))
        except ValueError:
            numeric_price = 0.0
        custom_date = request.form.get('updated_at')
        if custom_date:
            updated_at_str = f"{custom_date} {datetime.now().strftime('%H:%M:%S')}"
        else:
            updated_at_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        rtdb.reference('market_data').push({
            "commodity": request.form.get('commodity').strip(),
            "category": request.form.get('category', 'Other'), 
            "region": request.form.get('region'),
            "price": numeric_price,
            "unit": request.form.get('unit', 'kg'), 
            "currency": request.form.get('currency', 'USD'),
            "trend": request.form.get('trend'),
            "updated_at": updated_at_str 
        })

        flash(f"Market data for {request.form.get('commodity')} published successfully!", "success")
        return redirect(url_for('market_data_manager'))
        
    items = rtdb.reference('market_data').get() or {}
    
    # Clean up the data for the frontend
    market_list = []
    for k, v in items.items():
        # Handle legacy data that might be missing the new fields
        v['id'] = k
        v['category'] = v.get('category', 'General')
        v['unit'] = v.get('unit', 'unit')
        # Convert Firebase '{".sv": "timestamp"}' to a string if it exists in legacy data
        if isinstance(v.get('updated_at'), dict):
            v['updated_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        market_list.append(v)
        
    # Convert the iterator back into a proper list so JSON can read it!
    return render_template('market data manager.html', market_items=list(reversed(market_list)))


@app.route('/admin/delete-market-data/<item_id>', methods=['POST'])
@admin_required
def delete_market_data(item_id):
    try: 
        rtdb.reference(f'market_data/{item_id}').delete()
        flash("Removed.", "success")
    except Exception: 
        flash("Error.", "danger")
    return redirect(url_for('market_data_manager'))

@app.route('/admin/zim-data-manager', methods=['GET', 'POST'])
@admin_required
def zim_market_data_manager():
    if request.method == 'POST':
        try:
            numeric_price = float(request.form.get('price', 0))
        except ValueError:
            numeric_price = 0.0
        custom_date = request.form.get('updated_at')
        if custom_date:
            updated_at_str = f"{custom_date} {datetime.now().strftime('%H:%M:%S')}"
        else:
            updated_at_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        rtdb.reference('market_data').push({
            "commodity": request.form.get('commodity').strip(),
            "category": request.form.get('category', 'Other'), 
            "region": request.form.get('region'),
            "price": numeric_price,
            "unit": request.form.get('unit', 'kg'), 
            "currency": request.form.get('currency', 'USD'),
            "trend": request.form.get('trend'),
            "updated_at": updated_at_str,
            "country": "Zimbabwe"
        })

        flash(f"Zimbabwe Market data for {request.form.get('commodity')} published successfully!", "success")
        return redirect(url_for('zim_market_data_manager'))
        
    items = rtdb.reference('market_data').get() or {}
    market_list = []
    for k, v in items.items():
        if v.get('country') == 'Zimbabwe':
            v['id'] = k
            market_list.append(v)
        
    return render_template('zim_market_data_manager.html', market_items=list(reversed(market_list)))

@app.route('/admin/delete-zim-market-data/<item_id>', methods=['POST'])
@admin_required
def delete_zim_market_data(item_id):
    try: 
        rtdb.reference(f'market_data/{item_id}').delete()
        flash("Removed Zimbabwe market record.", "success")
    except Exception: 
        flash("Error.", "danger")
    return redirect(url_for('zim_market_data_manager'))

@app.route('/admin/content', methods=['GET', 'POST'])
@admin_required
def content_manager():
    cms_ref = rtdb.reference('site_content')
    history_ref = rtdb.reference('content_history')
    if request.method == 'POST':
        page_id = request.form.get('page_selection')
        title = request.form.get('content_title')
        body = request.form.get('body_text')
        cms_ref.child(page_id).set({'title': title, 'body': body, 'updated_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")})
        history_ref.push({'page': page_id, 'summary': f"Updated {title[:20]}...", 'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")})
        flash(f"Content for '{page_id}' successfully synchronized.", "success")
        return redirect(url_for('content_manager'))
    recent_edits = history_ref.order_by_key().limit_to_last(5).get()
    edits_list = list(recent_edits.values()) if recent_edits else []
    edits_list.reverse() 
    return render_template('content manager.html', recent_edits=edits_list)

@app.route('/admin/subscribers')
@admin_required
def subscriber_management():
    try:
        all_users = rtdb.reference('users').get()
        subscribers_list = [{'uid': uid, **data} for uid, data in all_users.items()] if all_users else []
        return render_template('subscriber management.html', subscribers=subscribers_list)
    except Exception as e:
        flash("Could not load the subscribers list.", "danger")
        return redirect(url_for('admin_dashboard'))
    
@app.route('/admin/update-role', methods=['POST'])
@admin_required
def update_user_role():
    """Updates a user's role (admin/tutor/client) and tier (free/premium)."""
    target_uid = request.form.get('user_id')
    new_role = request.form.get('role')
    new_tier = request.form.get('tier')

    if not target_uid:
        flash("User ID is missing.", "danger")
        return redirect(url_for('subscriber_management'))

    try:
        rtdb.reference(f'users/{target_uid}').update({
            'role': new_role,
            'subscription_tier': new_tier
        })
        flash(f"User updated successfully to {new_role} ({new_tier}).", "success")
    except Exception as e:
        flash(f"Error updating user: {e}", "danger")

    return redirect(url_for('subscriber_management'))

# ==========================================
# PUBLIC MARKET INTELLIGENCE & AI
# ==========================================
@app.route('/market-intelligence')
def market_intelligence():
    try:
        ref = rtdb.reference('market_data')
        all_items = ref.get() or {}
        
        # Automatic Kakamega Market Seed if empty or very few records
        if not all_items or len(all_items) < 5:
            seed_data = [
                # Maize
                {"commodity": "Maize (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 3200.0, "unit": "per 90kg bag", "currency": "KES", "trend": "stable", "updated_at": "2026-05-25 10:00:00"},
                {"commodity": "Maize (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 3100.0, "unit": "per 90kg bag", "currency": "KES", "trend": "down", "updated_at": "2026-05-26 10:00:00"},
                {"commodity": "Maize (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 3150.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-27 10:00:00"},
                {"commodity": "Maize (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 3000.0, "unit": "per 90kg bag", "currency": "KES", "trend": "down", "updated_at": "2026-05-28 10:00:00"},
                {"commodity": "Maize (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 2800.0, "unit": "per 90kg bag", "currency": "KES", "trend": "down", "updated_at": "2026-05-29 10:00:00"},
                
                # Beans
                {"commodity": "Rosecoco Beans (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 8500.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-25 10:00:00"},
                {"commodity": "Rosecoco Beans (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 8700.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-26 10:00:00"},
                {"commodity": "Rosecoco Beans (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 8600.0, "unit": "per 90kg bag", "currency": "KES", "trend": "down", "updated_at": "2026-05-27 10:00:00"},
                {"commodity": "Rosecoco Beans (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 8800.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-28 10:00:00"},
                {"commodity": "Rosecoco Beans (90kg)", "category": "Grains", "region": "Kakamega Central Market", "price": 9000.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-29 10:00:00"},

                # Sukuma Wiki
                {"commodity": "Sukuma Wiki (Crate)", "category": "Vegetables", "region": "Lurambi Market Kakamega", "price": 1200.0, "unit": "per large crate", "currency": "KES", "trend": "down", "updated_at": "2026-05-25 10:00:00"},
                {"commodity": "Sukuma Wiki (Crate)", "category": "Vegetables", "region": "Lurambi Market Kakamega", "price": 1100.0, "unit": "per large crate", "currency": "KES", "trend": "down", "updated_at": "2026-05-26 10:00:00"},
                {"commodity": "Sukuma Wiki (Crate)", "category": "Vegetables", "region": "Lurambi Market Kakamega", "price": 1150.0, "unit": "per large crate", "currency": "KES", "trend": "up", "updated_at": "2026-05-27 10:00:00"},
                {"commodity": "Sukuma Wiki (Crate)", "category": "Vegetables", "region": "Lurambi Market Kakamega", "price": 1050.0, "unit": "per large crate", "currency": "KES", "trend": "down", "updated_at": "2026-05-28 10:00:00"},
                {"commodity": "Sukuma Wiki (Crate)", "category": "Vegetables", "region": "Lurambi Market Kakamega", "price": 1000.0, "unit": "per large crate", "currency": "KES", "trend": "down", "updated_at": "2026-05-29 10:00:00"},

                # Sweet Potatoes
                {"commodity": "Sweet Potatoes (Bag)", "category": "Tubers", "region": "Mumias Market Kakamega", "price": 4200.0, "unit": "per 90kg bag", "currency": "KES", "trend": "stable", "updated_at": "2026-05-25 10:00:00"},
                {"commodity": "Sweet Potatoes (Bag)", "category": "Tubers", "region": "Mumias Market Kakamega", "price": 4300.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-26 10:00:00"},
                {"commodity": "Sweet Potatoes (Bag)", "category": "Tubers", "region": "Mumias Market Kakamega", "price": 4150.0, "unit": "per 90kg bag", "currency": "KES", "trend": "down", "updated_at": "2026-05-27 10:00:00"},
                {"commodity": "Sweet Potatoes (Bag)", "category": "Tubers", "region": "Mumias Market Kakamega", "price": 4400.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-28 10:00:00"},
                {"commodity": "Sweet Potatoes (Bag)", "category": "Tubers", "region": "Mumias Market Kakamega", "price": 4500.0, "unit": "per 90kg bag", "currency": "KES", "trend": "up", "updated_at": "2026-05-29 10:00:00"},
            ]
            for item in seed_data:
                ref.push(item)
            all_items = ref.get() or {}

        # Form list of dictionaries with ID and clean date conversions
        full_list = []
        for k, v in all_items.items():
            v['id'] = k
            if isinstance(v.get('updated_at'), dict):
                v['updated_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            full_list.append(v)

        # Get latest unique commodity entries for the preview list
        unique_commodities = {}
        sorted_full = sorted(full_list, key=lambda x: x.get('updated_at', ''))
        for item in sorted_full:
            unique_commodities[item['commodity']] = item
        
        # Convert unique dict to preview_list
        preview_list = list(unique_commodities.values())
        # Sort by updated_at descending for latest first
        preview_list.sort(key=lambda x: x.get('updated_at', ''), reverse=True)

        return render_template('market intelligence.html', preview=preview_list, all_market_items=full_list)
    except Exception as e:
        print(f"Error in market intelligence: {e}")
        return render_template('market intelligence.html', preview=[], all_market_items=[])

@app.route('/live-market-prices')
@login_required
@premium_required
def live_market_prices():
    items = MarketData.query.order_by(MarketData.commodity.asc()).all()
    return render_template('live market prices.html', market_items=items)

@app.route('/trends-forecasts')
@login_required
@premium_required 
def trends_forecasts():
    records = MarketData.query.filter_by(commodity="Maize (90kg)").all()
    hist = [{'date': r.updated_at, 'price': r.price} for r in records]
    labels = [r.updated_at.strftime('%b %d') for r in records]
    prices = [r.price for r in records]
    ai = generate_price_forecast(hist, 5) if len(hist) >= 5 else {}
    if ai and "error" not in ai:
        labels.extend(ai['future_dates']); prices.extend(ai['predicted_prices'])
    return render_template('trends&forecasts.html', labels=labels, prices=prices, ai_insight=ai)

@app.route('/api/market-prices', methods=['GET'])
def api_market_prices():
    if session.get('tier') not in ['premium', 'pro'] and session.get('user_role') not in ['admin', 'tutor']:
        return jsonify({"error": "Premium subscription required to access raw data"}), 403
        
    items = rtdb.reference('market_data').get() or {}
    return jsonify([{'id': k, **v} for k, v in items.items()]), 200

# ==========================================
# FARMERMAN ACADEMY (STUDENT ROUTES)
# ==========================================
@app.route('/academy')
@login_required 
def academy_home():
    user_tier = session.get('subscription_tier', 'free')
    if session.get('role') == 'admin':
        user_tier = 'admin' 
        
    try:
        courses_data = rtdb.reference('academy_courses').get() or {}
        courses = [{'id': k, **v} for k, v in courses_data.items()]
        courses.reverse() 
    except Exception as e:
        print(f"Error fetching courses: {e}")
        courses = []
        
    return render_template('academy/index.html', user_tier=user_tier, courses=courses)

@app.route('/academy/my-learning')
@login_required
def my_learning():
    user_id = session.get('user_id')
    user_tier = session.get('subscription_tier', 'free')
    
    all_courses_data = rtdb.reference('academy_courses').get() or {}
    user_progress_data = rtdb.reference(f'user_progress/{user_id}').get() or {}
    
    enrolled_courses = []
    completed_courses = []
    
    for cid, cdata in all_courses_data.items():
        cdata['id'] = cid
        
        if cid in user_progress_data:
            progress = int(user_progress_data[cid].get('progress', 0))
            cdata['progress'] = progress
            
            if progress >= 100:
                completed_courses.append(cdata)
            else:
                enrolled_courses.append(cdata)
        else:
            cdata['progress'] = 0 
            enrolled_courses.append(cdata)

    return render_template('academy/my_learning.html', 
                           enrolled_courses=enrolled_courses,
                           completed_courses=completed_courses,
                           user_tier=user_tier)

@app.route('/academy/course/<course_id>')
@premium_required 
def view_lesson(course_id):
    lesson_data = rtdb.reference(f'academy_courses/{course_id}').get()
    
    if not lesson_data:
        flash("This course could not be found.", "warning")
        return redirect(url_for('academy_home'))
        
    lesson_data['id'] = course_id 
    
    comments_data = rtdb.reference(f'course_comments/{course_id}').get() or {}
    comments = [{'id': k, **v} for k, v in comments_data.items()]
    comments.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
    
    return render_template('academy/course_view.html', lesson=lesson_data, comments=comments)

@app.route('/academy/course/<course_id>/quiz')
@premium_required
def take_quiz(course_id):
    lesson_data = rtdb.reference(f'academy_courses/{course_id}').get()
    if not lesson_data:
        flash("Course not found.", "warning")
        return redirect(url_for('academy_home'))
        
    quiz_data = rtdb.reference(f'academy_courses/{course_id}/quiz').get()
    
    if not quiz_data:
        quiz_data = [
            {
                "question": "What is the optimal soil pH range for most agricultural crops?",
                "options": ["4.0 - 5.0 (Highly Acidic)", "6.0 - 7.0 (Slightly Acidic to Neutral)", "8.0 - 9.0 (Highly Alkaline)", "It does not matter"],
                "answer": 1 
            },
            {
                "question": "Which primary macronutrient is directly responsible for vigorous, leafy green growth?",
                "options": ["Phosphorus (P)", "Potassium (K)", "Nitrogen (N)", "Calcium (Ca)"],
                "answer": 2 
            },
            {
                "question": "When is the best time of day to apply foliar fertilizers or pesticides?",
                "options": ["Early morning or late afternoon", "High noon when the sun is brightest", "During a heavy rainstorm", "Midnight"],
                "answer": 0 
            }
        ]
        
    return render_template('academy/quiz_view.html', 
                           lesson=lesson_data, 
                           quiz=quiz_data, 
                           course_id=course_id)

@app.route('/academy/certificate')
@login_required
def generate_certificate():
    user_id = session.get('user_id')
    user_profile = rtdb.reference(f'users/{user_id}').get() or {}
    full_name = user_profile.get('full_name', 'Esteemed Farmer')
    
    user_progress_data = rtdb.reference(f'user_progress/{user_id}').get() or {}
    has_completed_course = any(int(data.get('progress', 0)) >= 100 for data in user_progress_data.values())
            
    if not has_completed_course and session.get('role') != 'admin':
        flash("You must complete at least one course to earn a certificate.", "warning")
        return redirect(url_for('my_learning'))
        
    today_date = datetime.now().strftime("%B %d, %Y")
    return render_template('academy/certificate.html', student_name=full_name, date=today_date)

@app.route('/academy/leaderboard')
@login_required
def academy_leaderboard():
    all_users = rtdb.reference('users').get() or {}
    all_progress = rtdb.reference('user_progress').get() or {}
    leaderboard = []
    
    for uid, user_data in all_users.items():
        # Allow legacy clients, new buyers, and new sellers to appear on the board
        user_role = user_data.get('role', 'buyer').lower()
        if user_role not in ['client', 'buyer', 'seller']:
            continue
            
        name = user_data.get('full_name', 'Anonymous Farmer')
        points = sum(100 for course_data in all_progress.get(uid, {}).values() if int(course_data.get('progress', 0)) >= 100)
                    
        if points > 0:
            leaderboard.append({'uid': uid, 'name': name, 'points': points})
            
    leaderboard.sort(key=lambda x: x['points'], reverse=True)
    
    current_user_id = session.get('user_id')
    current_user_rank = next(({'rank': i + 1, 'points': u['points']} for i, u in enumerate(leaderboard) if u['uid'] == current_user_id), None)

    return render_template('academy/academy_leaderboard.html', 
                           leaderboard=leaderboard, 
                           current_user_rank=current_user_rank) 

@app.route('/agripreneur_training')
@premium_required
def agripreneur_training():
    try:
        content_ref = rtdb.reference('training_content').get() or {}
        all_content = sorted(content_ref.values(), key=lambda x: x.get('order', 0))
    except Exception as e:
        print(f"Academy Data Error: {e}")
        all_content = []

    categories = ['agripreneur', 'aqua', 'econ']
    categorized_content = {cat: [c for c in all_content if c.get('category') == cat] for cat in categories}
    
    return render_template('agripreneur_training.html', content=categorized_content)

@app.route('/secure-media/<path:filename>')
@premium_required
def secure_media(filename):
    try:
        response = send_from_directory(PREMIUM_CONTENT_FOLDER, filename)
        response.headers['X-Content-Type-Options'] = 'nosniff'
        return response
    except FileNotFoundError:
        abort(404, description="This training resource is currently unavailable.")
        
@app.route('/admin-delete-training/<content_id>')
@login_required
def delete_training(content_id):
    ref = rtdb.reference(f'training_content/{content_id}')
    item = ref.get()
    
    if item:
        try:
            os.remove(os.path.join(PREMIUM_CONTENT_FOLDER, item['filename']))
        except OSError:
            pass 
        
        ref.delete()
        flash("Content removed successfully.", "info")
    
    return redirect(url_for('training_manager'))

# ==========================================
# FARMERMAN ACADEMY (API ROUTES)
# ==========================================
@app.route('/api/academy/update-progress', methods=['POST'])
@login_required
def update_progress():
    data = request.json
    course_id = data.get('course_id')
    progress = data.get('progress', 0)
    user_id = session.get('user_id')
    
    if not course_id:
        return jsonify({"error": "Missing course ID"}), 400
        
    try:
        rtdb.reference(f'user_progress/{user_id}/{course_id}').update({
            'progress': progress,
            'last_accessed': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        return jsonify({"success": True, "message": "Progress updated"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/academy/submit-quiz', methods=['POST'])
@login_required
def submit_quiz():
    data = request.json
    course_id = data.get('course_id')
    score = data.get('score', 0)
    user_id = session.get('user_id')
    
    if not course_id:
        return jsonify({"error": "Missing course ID"}), 400
        
    try:
        rtdb.reference(f'user_progress/{user_id}/{course_id}').update({
            'quiz_score': score,
            'quiz_completed_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        return jsonify({"success": True, "message": "Score saved successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/academy/course/<course_id>/comment', methods=['POST'])
@premium_required
def post_course_comment(course_id):
    message = request.form.get('message')
    user_id = session.get('user_id')
    
    if message and message.strip():
        user_profile = rtdb.reference(f'users/{user_id}').get() or {}
        full_name = user_profile.get('full_name', 'FarmerMan Scholar')
        
        rtdb.reference(f'course_comments/{course_id}').push({
            'user_id': user_id,
            'user_name': full_name,
            'message': message.strip(),
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        flash("Your question has been posted to the discussion!", "success")
    else:
        flash("Comment cannot be empty.", "warning")
        
    return redirect(url_for('view_lesson', course_id=course_id))

# ==========================================
# FARMERMAN ACADEMY (TUTOR ROUTES)
# ==========================================
@app.route('/academy/tutor/dashboard')
@tutor_required
def tutor_dashboard():
    tutor_id = session.get('user_id')
    
    all_courses = rtdb.reference('academy_courses').get() or {}
    my_courses = []
    
    for cid, cdata in all_courses.items():
        if cdata.get('tutor_id') == tutor_id or session.get('role') == 'admin':
            cdata['id'] = cid
            my_courses.append(cdata)
            
    all_progress = rtdb.reference('user_progress').get() or {}
    total_enrollments = 0
    total_completions = 0
    
    my_course_ids = [c['id'] for c in my_courses]
    
    for uid, progress_data in all_progress.items():
        for cid, data in progress_data.items():
            if cid in my_course_ids:
                total_enrollments += 1
                if int(data.get('progress', 0)) >= 100:
                    total_completions += 1

    return render_template('academy/tutor_dashboard.html', 
                           my_courses=my_courses,
                           total_enrollments=total_enrollments,
                           total_completions=total_completions)

@app.route('/academy/tutor/gradebook')
@tutor_required
def tutor_gradebook():
    all_users = rtdb.reference('users').get() or {}
    all_courses = rtdb.reference('academy_courses').get() or {}
    all_progress = rtdb.reference('user_progress').get() or {}
    
    student_records = []
    total_completions = 0
    
    for uid, progress_data in all_progress.items():
        user_info = all_users.get(uid, {})
        student_name = user_info.get('full_name', 'Unknown Farmer')
        student_email = user_info.get('email', 'N/A')
        
        for course_id, data in progress_data.items():
            course_info = all_courses.get(course_id, {})
            course_title = course_info.get('title', 'Deleted Course')
            progress_score = int(data.get('progress', 0))
            quiz_score = data.get('quiz_score')
            last_accessed = data.get('last_accessed', 'Unknown')
            
            if progress_score >= 100:
                total_completions += 1
                
            student_records.append({
                'student_name': student_name,
                'student_email': student_email,
                'course_title': course_title,
                'progress': progress_score,
                'quiz_score': quiz_score if quiz_score is not None else 'Not Taken',
                'last_accessed': last_accessed[:10] if last_accessed != 'Unknown' else 'N/A'
            })
            
    student_records.sort(key=lambda x: x['last_accessed'], reverse=True)
    active_learners = len(all_progress.keys())
    
    return render_template('academy/gradebook.html', 
                           student_records=student_records,
                           active_learners=active_learners,
                           total_completions=total_completions)

@app.route('/academy/tutor/sessions', methods=['GET', 'POST'])
@tutor_required
def tutor_sessions():
    tutor_id = session.get('user_id')
    
    if request.method == 'POST':
        course_id = request.form.get('course_id')
        topic = request.form.get('topic')
        session_date = request.form.get('session_date')
        session_time = request.form.get('session_time')
        meet_link = request.form.get('meet_link')
        
        try:
            rtdb.reference('tutor_sessions').push({
                'tutor_id': tutor_id,
                'course_id': course_id,
                'topic': topic,
                'date': session_date,
                'time': session_time,
                'meet_link': meet_link,
                'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            
            if course_id:
                rtdb.reference(f'academy_courses/{course_id}').update({'meet_link': meet_link})
                
            flash("Live session scheduled successfully! The course page has been updated.", "success")
            return redirect(url_for('tutor_sessions'))
            
        except Exception as e:
            flash(f"Error scheduling session: {e}", "danger")

    all_courses = rtdb.reference('academy_courses').get() or {}
    my_courses = [{'id': k, 'title': v.get('title')} for k, v in all_courses.items() 
                  if v.get('tutor_id') == tutor_id or session.get('role') == 'admin']
    
    all_sessions = rtdb.reference('tutor_sessions').get() or {}
    my_sessions = []
    
    for sid, sdata in all_sessions.items():
        if sdata.get('tutor_id') == tutor_id or session.get('role') == 'admin':
            course = all_courses.get(sdata.get('course_id'), {})
            sdata['course_title'] = course.get('title', 'Deleted Course')
            sdata['id'] = sid
            my_sessions.append(sdata)
            
    my_sessions.sort(key=lambda x: f"{x.get('date')} {x.get('time')}", reverse=True)
    
    return render_template('academy/tutor_sessions.html', 
                           my_courses=my_courses, 
                           my_sessions=my_sessions)
          
# ==========================================
# TUTOR SECURITY & UPLOAD LOGIC
# ==========================================
@app.route('/academy/tutor/builder', methods=['GET', 'POST'])
@tutor_required
def course_builder():
    if request.method == 'POST':
        # 1. Grab text data from the form
        title = request.form.get('course_title')
        description = request.form.get('course_description')
        category = request.form.get('category')
        meet_link = request.form.get('meet_link', '')       
        # 2. Handle File Uploads (Video & PDF)
        video_file = request.files.get('video_file')
        resource_file = request.files.get('resource_file')
        
        video_filename = ""
        resource_filename = ""
        
        # Save Video
        if video_file and allowed_training_file(video_file.filename):
            video_filename = secure_filename(video_file.filename)
            video_file.save(os.path.join(PREMIUM_CONTENT_FOLDER, video_filename))
            
        # Save PDF Resource
        if resource_file and allowed_training_file(resource_file.filename):
            resource_filename = secure_filename(resource_file.filename)
            resource_file.save(os.path.join(PREMIUM_CONTENT_FOLDER, resource_filename))
            
        # 3. Save everything to Firebase Realtime Database
        try:
            rtdb.reference('academy_courses').push({
                'title': title,
                'description': description,
                'category': category,
                'meet_link': meet_link,
                'video_file': video_filename,
                'resource_file': resource_filename,
                'tutor_id': session.get('user_id'),
                'tutor_name': session.get('user_email'), # Or fetch their full name
                'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            flash(f"Success! '{title}' has been published to the Academy.", "success")
            return redirect(url_for('course_builder'))
        except Exception as e:
            flash(f"Database Error: {str(e)}", "danger")

    return render_template('academy/course_builder.html')

# ==========================================
# PAYMENTS & CALLBACKS
# ==========================================
stripe.api_key = os.environ.get('STRIPE_SECRET_KEY')
PAYPAL_CLIENT_ID = os.environ.get('PAYPAL_CLIENT_ID')
PAYSTACK_SECRET_KEY = os.environ.get('PAYSTACK_SECRET_KEY')

# --- MASTER PRICING DICTIONARY ---
# Single source of truth for all payment gateways
SYSTEM_PRICING = {
    'seed': {"name": "Seed Package", "kes": 450, "usd": 3.00, "frequency": "1 update/week"},
    'growth': {"name": "Growth Package", "kes": 750, "usd": 5.00, "frequency": "2 updates/week"},
    'harvest': {"name": "Harvest Package", "kes": 1500, "usd": 10.00, "frequency": "3 updates/week"},
    'market_intel': {"name": "Market Intelligence Monthly Standalone", "kes": 50, "usd": 0.50},
    'bronze': {"name": "Bronze Consulting Package", "kes": 1000, "usd": 10.00},
    'silver': {"name": "Silver Consulting Package", "kes": 2000, "usd": 20.00},
    'gold': {"name": "Gold Consulting Package", "kes": 3000, "usd": 30.00},
    'course_entrepreneurship': {"name": "Agribusiness Entrepreneurship", "kes": 3750, "usd": 25.00},
    'course_marketing': {"name": "Agricultural Marketing", "kes": 3000, "usd": 20.00},
    'course_finance': {"name": "Financial Management", "kes": 4500, "usd": 30.00},
    'course_value_chain': {"name": "Value Chain Development", "kes": 5250, "usd": 35.00},
    'course_post_harvest': {"name": "Post-Harvest Loss Strategies", "kes": 3000, "usd": 20.00},
    'course_bankable': {"name": "Bankable Agribusiness Projects", "kes": 6000, "usd": 40.00},
    'course_market_intel': {"name": "Market Intelligence & Research", "kes": 7500, "usd": 50.00},
    'course_consulting': {"name": "Agribusiness Consulting Skills", "kes": 9000, "usd": 60.00},
    'bundle_starter': {"name": "Agribusiness Starter Bundle", "kes": 9000, "usd": 60.00},
    'bundle_investor': {"name": "Agribusiness Investor Bundle", "kes": 18000, "usd": 120.00}
}

def get_plan_price(plan_id, category=None):
    """
    Returns Kes & Usd values. For market_intel, supports:
    individual (50 KSH / $0.50 monthly), cbo (300 KSH / $3.00 monthly),
    company (900 KSH / $9.00 monthly), ngo / other (2700 KSH / $27.00 monthly).
    """
    base_plan = SYSTEM_PRICING.get(plan_id, SYSTEM_PRICING['bronze'])
    if plan_id == 'market_intel':
        if not category:
            category = 'individual'
        category = str(category).lower().strip()
        if category == 'individual':
            return {"name": "Market Intelligence - Individual (Monthly)", "kes": 50, "usd": 0.50}
        elif category == 'cbo':
            return {"name": "Market Intelligence - CBO (Monthly)", "kes": 300, "usd": 3.00}
        elif category == 'company':
            return {"name": "Market Intelligence - Company (Monthly)", "kes": 900, "usd": 9.00}
        elif category in ['ngo', 'other']:
            return {"name": "Market Intelligence - NGO/Other (Monthly)", "kes": 2700, "usd": 27.00}
    return {
        "name": base_plan['name'],
        "kes": base_plan['kes'],
        "usd": base_plan['usd']
    }

def record_successful_transaction(user_id, plan_id, amount, gateway, receipt_number):
    """Updates the user tier, status, and expiry date. Logs the transaction."""
    try:
        # 1. Calculate Expiry (30 days from now)
        eat_tz = timezone(timedelta(hours=3))
        now = datetime.now(eat_tz)
        expiry_date = now + timedelta(days=30)
        
        # 2. Upgrade the user's tier if they are a registered user
        if user_id and user_id != 'guest':
            rtdb.reference(f'users/{user_id}').update({
                'subscription_tier': plan_id,
                'subscription_status': 'active',
                'subscription_expiry': expiry_date.strftime("%Y-%m-%d %H:%M:%S")
            })
            
            # Sync with SQLAlchemy for local analytics if needed
            local_user = sqlalchemy_db.session.get(User, user_id) if isinstance(user_id, int) else User.query.filter_by(email=user_id).first()
            if local_user:
                local_user.subscription_tier = plan_id
                local_user.subscription_status = 'active'
                local_user.subscription_expiry = expiry_date
                sqlalchemy_db.session.commit()
        
        # 3. Update active session
        try:
            if session.get('user_id') == user_id:
                session['tier'] = plan_id
                session['subscription_tier'] = plan_id
                session['subscription_status'] = 'active'
            
            if plan_id in ['seed', 'growth', 'harvest', 'market_intel'] and (not user_id or user_id == 'guest'):
                session['market_intel_guest'] = True
        except Exception:
            pass
            
        # 4. Log the transaction
        log_id = user_id if user_id and user_id != 'guest' else f"guest_{receipt_number}"
        rtdb.reference(f'completed_transactions/{log_id}').push({
            'receipt_number': receipt_number, 
            'amount': amount,
            'gateway': gateway,
            'plan_purchased': plan_id,
            'date': now.strftime("%Y-%m-%d %H:%M:%S")
        })
        return True
    except Exception as e:
        print(f"CRITICAL: Transaction Sync Error: {e}")
        return False

# --- 1. M-PESA LOGIC ---
@app.route('/process-mpesa', methods=['POST'])
def process_mpesa():
    phone = request.form.get('phone_number')
    plan_id = request.form.get('plan_id', 'bronze') 
    raw_amount = request.form.get('amount') 
    category = request.form.get('category')
    
    # Store plan in session for guest tracking
    session['pending_plan'] = plan_id
    
    # Validation & Fallbacks - Now uses get_plan_price
    try:
        amount = int(float(raw_amount)) if raw_amount else get_plan_price(plan_id, category)['kes']
    except (ValueError, TypeError):
        amount = get_plan_price(plan_id, category)['kes']
        
    try:
        res = initiate_stk_push(phone, amount)
        if res and res.get('ResponseCode') == '0':
            checkout_id = res.get("CheckoutRequestID")
            # Store pending transaction for callback matching
            rtdb.reference(f'pending_transactions/{checkout_id}').set({
                'user_id': session.get('user_id', 'guest'), 
                'amount': amount, 
                'plan_id': plan_id, 
                'category': category or '',
                'status': 'awaiting_payment',
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            
            return redirect(url_for('payment_processing', checkout_id=checkout_id))
            
        error_msg = res.get('errorMessage', 'Safaricom service is currently unavailable.')
        return redirect(url_for('payment_failed', msg=f"M-Pesa Error: {error_msg}"))
        
    except Exception as e:
        print(f"CRITICAL MPESA STK ERROR: {str(e)}")
        return redirect(url_for('payment_failed', msg="M-Pesa Gateway is currently unstable. Please try again later or use a Card."))

    
@app.route('/mpesa-callback', methods=['POST'])
def mpesa_callback():
    """Receives async confirmation from Safaricom for BOTH Subscriptions & Banking."""
    try:
        data = request.json
        stk = data.get('Body', {}).get('stkCallback', {})
        checkout_id = stk.get("CheckoutRequestID")
        
        pending_ref = rtdb.reference(f'pending_transactions/{checkout_id}')
        pending_data = pending_ref.get()

        if pending_data:
            # Define timezone for accurate logging
            eat_tz = timezone(timedelta(hours=3))
            
            # Fetch the tx_type. Default to 'subscription' for older/legacy API calls
            tx_type = pending_data.get('tx_type', 'subscription')
            
            # 1. SUCCESS: Safaricom confirms the PIN was entered and funds captured
            if stk.get('ResultCode') == 0:
                meta = stk.get('CallbackMetadata', {}).get('Item', [])
                receipt = next((i['Value'] for i in meta if i['Name'] == 'MpesaReceiptNumber'), 'UNKNOWN')
                
                uid = pending_data.get('user_id')
                amount = float(pending_data.get('amount', 0)) # Ensure amount is a float for math
                
                if tx_type == 'banking_deposit':
                    # --- ROUTE TO BANKING SYSTEM ---
                    fund_type = pending_data.get('fund_type')
                    
                    if fund_type == 'emergency_fund':
                        account_ref = rtdb.reference(f'banking_accounts/{uid}')
                        current_bal = float(account_ref.child('emergency_fund').get() or 0.0)
                        account_ref.update({'emergency_fund': current_bal + amount})
                    else:
                        # It is a specific group ID
                        account_ref = rtdb.reference(f'banking_accounts/{uid}/standard_savings/{fund_type}')
                        current_bal = float(account_ref.get() or 0.0)
                        account_ref.set(current_bal + amount)
                    
                    # Log the banking transaction
                    rtdb.reference(f'banking_transactions/{uid}').push({
                        'type': 'deposit',
                        'fund_type': fund_type,
                        'amount': amount,
                        'receipt': receipt,
                        'timestamp': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S")
                    })
                else:
                    # --- ROUTE TO SUBSCRIPTION SYSTEM ---
                    plan_bought = pending_data.get('plan_id', 'bronze')
                    category = pending_data.get('category')
                    record_successful_transaction(uid, plan_bought, amount, "M-Pesa", receipt)
                    if plan_bought == 'market_intel' and category and uid != 'guest':
                        try:
                            rtdb.reference(f'users/{uid}').update({'organization_category': category})
                        except Exception as ex:
                            print(f"Error saving category to database: {ex}")
                    
                # Cleanup deletes the pending record
                pending_ref.delete() 
            
            # 2. FAILED: User cancelled, typed wrong PIN, or had insufficient funds
            else:
                # Update status so the frontend polling knows to show the failed screen
                pending_ref.update({'status': 'failed'})
                
        # Always return 200 OK so Safaricom doesn't keep retrying
        return jsonify({"ResultCode": 0, "ResultDesc": "Accepted"}), 200
        
    except Exception as e:
        print(f"M-Pesa Callback Error: {e}")
        return jsonify({"ResultCode": 1, "ResultDesc": "Internal Server Error"}), 500
    
    
@app.route('/checkout')
def subscriber_checkout():
    # 1. Get the plan from the URL (e.g., /checkout?plan=silver or ?plan=course_marketing)
    # If no plan is specified, it safely defaults to 'bronze'
    plan_id = request.args.get('plan', 'bronze')
    
    # Store plan in session
    session['pending_plan'] = plan_id
    
    # 2. Look up the exact pricing from our master dictionary (SYSTEM_PRICING)
    selected_plan = SYSTEM_PRICING.get(plan_id, SYSTEM_PRICING['bronze'])

    # 3. Send the dynamic prices to the HTML template
    return render_template(
        'payments/subscriber_checkout.html',
        plan_id=plan_id, 
        plan_name=selected_plan['name'], 
        amount_kes=selected_plan['kes'], 
        amount_usd=selected_plan['usd'],
        paystack_public_key=os.environ.get('PAYSTACK_PUBLIC_KEY', ''),
        stripe_public_key=os.environ.get('STRIPE_PUBLIC_KEY', ''),
        paypal_client_id=os.environ.get('PAYPAL_CLIENT_ID', '')
    )

@app.route('/payment-processing/<checkout_id>')
def payment_processing(checkout_id):
    """Renders the waiting room while the user types their M-Pesa PIN."""
    return render_template('payments/payment_processing.html', checkout_id=checkout_id)

@app.route('/api/check-payment/<checkout_id>')
def check_payment_status(checkout_id):
    """The frontend polls this endpoint every 3 seconds to check for the callback."""
    pending_txn = rtdb.reference(f'pending_transactions/{checkout_id}').get()
    
    if pending_txn:
        if pending_txn.get('status') == 'failed':
            return jsonify({'status': 'failed'})
        return jsonify({'status': 'pending'})
    
    # If the transaction is GONE, mpesa_callback successfully processed it!
    # If it was a market_intel plan, grant guest access
    if session.get('pending_plan') == 'market_intel':
        session['market_intel_guest'] = True
        
    return jsonify({'status': 'completed'})

# --- 2. STRIPE LOGIC ---
@app.route('/create-stripe-session', methods=['POST'])
def create_stripe_session():
    try:
        data = request.json
        plan_id = data.get('plan', 'bronze')
        category = data.get('category')
        
        # Store plan in session
        session['pending_plan'] = plan_id

        # Look up the plan using our dynamic get_plan_price function
        plan_details = get_plan_price(plan_id, category)
        # Stripe expects amounts in cents, so we multiply the USD amount by 100
        amount_in_cents = int(plan_details['usd'] * 100)

        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price_data': {
                    'currency': 'usd',
                    'unit_amount': amount_in_cents,
                    'product_data': {
                        'name': f"Farmerman Systems: {plan_details['name']}",
                        'description': f"Payment for {plan_details['name']}"
                    },
                },
                'quantity': 1,
            }],
            mode='payment',
            # Metadata is crucial for tracking who paid what in the Stripe Dashboard
            metadata={
                'user_id': session.get('user_id', 'guest'), 
                'plan_id': plan_id,
                'category': category or ''
            }, 
            success_url=url_for('stripe_success', plan_id=plan_id, category=category or '', _external=True) + '&session_id={CHECKOUT_SESSION_ID}',
            cancel_url=url_for('pricing', _external=True),
        )
        return jsonify({'id': checkout_session.id})
    except Exception as e:
        print(f"Stripe Session Error: {e}")
        return jsonify(error="Unable to connect to Stripe."), 403

@app.route('/stripe-success')
def stripe_success():
    """Validates the redirect back from Stripe."""
    plan_id = request.args.get('plan_id', 'bronze')
    category = request.args.get('category')
    session_id = request.args.get('session_id')
    user_id = session.get('user_id', 'guest')
    
    # Grab the accurate USD price for the receipt
    plan_details = get_plan_price(plan_id, category)
    amount_paid = plan_details['usd']
    
    if session_id:
        record_successful_transaction(user_id, plan_id, amount_paid, "Stripe", f"STR_{session_id[-8:]}")
        if plan_id == 'market_intel' and category and user_id != 'guest':
            try:
                rtdb.reference(f'users/{user_id}').update({'organization_category': category})
            except Exception as ex:
                print(f"Error saving category to database: {ex}")
        
        if plan_id == 'market_intel':
            session['market_intel_guest'] = True
            
        return redirect(url_for('payment_success'))
    
    return redirect(url_for('pricing'))

# --- 3. PAYPAL LOGIC ---
@app.route('/paypal-transaction-complete', methods=['POST'])
def paypal_transaction_complete():
    try:
        data = request.json
        order_id = data.get('orderID')
        plan_id = data.get('plan', 'bronze')
        category = data.get('category')
        user_id = session.get('user_id', 'guest')
        
        # Grab the accurate USD price using get_plan_price
        plan_details = get_plan_price(plan_id, category)
        amount_paid = plan_details['usd']
        
        if order_id:
            record_successful_transaction(user_id, plan_id, amount_paid, "PayPal", f"PAY_{order_id}")
            if plan_id == 'market_intel' and category and user_id != 'guest':
                try:
                    rtdb.reference(f'users/{user_id}').update({'organization_category': category})
                except Exception as ex:
                    print(f"Error saving category to database: {ex}")
            
            if plan_id == 'market_intel':
                session['market_intel_guest'] = True
                
            return jsonify({"status": "success"}), 200
            
        return jsonify({"status": "failed", "error": "Missing data"}), 400
    except Exception as e:
        print(f"PayPal Recording Error: {e}")
        return jsonify({"status": "failed"}), 500

# --- 4. PAYSTACK LOGIC ---
@app.route('/verify-paystack')
def verify_paystack():
    reference = request.args.get('reference')
    plan_id = request.args.get('plan', 'bronze')
    category = request.args.get('category')
    user_id = session.get('user_id', 'guest')
    
    if not reference: 
        flash("Invalid transaction reference.", "warning")
        return redirect(url_for('pricing'))
    
    secret_key = os.environ.get('PAYSTACK_SECRET_KEY')
    verify_url = f"https://api.paystack.co/transaction/verify/{reference}"
    headers = {"Authorization": f"Bearer {secret_key}"}
    
    try:
        response = requests.get(verify_url, headers=headers)
        response_data = response.json()
        
        if response_data.get('status') is True and response_data.get('data', {}).get('status') == 'success':
            # Paystack sends amount in Kobo/Cents, so we divide by 100
            actual_amount = response_data['data']['amount'] / 100 
            
            record_successful_transaction(user_id, plan_id, actual_amount, "Paystack", reference)
            if plan_id == 'market_intel' and category and user_id != 'guest':
                try:
                    rtdb.reference(f'users/{user_id}').update({'organization_category': category})
                except Exception as ex:
                    print(f"Error saving category to database: {ex}")
            
            if plan_id == 'market_intel':
                session['market_intel_guest'] = True
                
            flash(f"Payment successful! Welcome to the {plan_id.capitalize()} plan.", "success")
            return redirect(url_for('payment_success'))
            
        flash("Payment verification failed. Please contact support.", "danger")
    except Exception as e:
        print(f"Paystack Verification Error: {e}")
        flash("Server error during verification.", "danger")
        
    return redirect(url_for('pricing'))

def get_pesapal_ipn_id(token, host_url):
    """Retrieves or registers the IPN ID dynamically, caching it in Firebase for speed."""
    ipn_cache_key = host_url.replace('.', '_').replace('/', '_').replace(':', '_')
    try:
        cached = rtdb.reference(f'pesapal_ipns/{ipn_cache_key}').get()
        if cached and isinstance(cached, dict) and 'ipn_id' in cached:
            return cached['ipn_id']
    except Exception as ex:
        print(f"Error accessing firebase IPN cache: {ex}")
        
    # Register new IPN URL
    ipn_endpoint = f"{host_url}/pesapal-ipn"
    res = register_pesapal_ipn(token, ipn_endpoint)
    if res and isinstance(res, dict):
        ipn_id = res.get('ipn_id')
        if ipn_id:
            try:
                rtdb.reference(f'pesapal_ipns/{ipn_cache_key}').set({
                    'ipn_id': ipn_id,
                    'url': ipn_endpoint,
                    'registered_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
            except Exception as ex:
                print(f"Error caching IPN: {ex}")
            return ipn_id
    return None

# --- 5. PESAPAL LOGIC ---
@app.route('/create-pesapal-session', methods=['POST'])
def create_pesapal_session():
    try:
        data = request.json or {}
        plan_id = data.get('plan', 'bronze')
        category = data.get('category')
        
        plan_details = get_plan_price(plan_id, category)
        amount = float(plan_details['kes'])
        currency = "KES"
        
        user_id = session.get('user_id', 'guest')
        # Generate merchant reference
        merchant_ref = f"PESA_{user_id[-6:]}_{int(time.time())}"
        
        # Save pending transaction state
        rtdb.reference(f'pending_transactions/{merchant_ref}').set({
            'user_id': user_id,
            'amount': amount,
            'plan_id': plan_id,
            'category': category or '',
            'status': 'awaiting_payment',
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Determine host url dynamically
        host_url = request.host_url.rstrip('/')
        callback_url = f"{host_url}/pesapal-callback"
        
        # Authenticate and obtain token
        token = get_pesapal_token()
        if not token:
            return jsonify({"error": "Failed to authenticate with PesaPal API"}), 500
            
        ipn_id = get_pesapal_ipn_id(token, host_url)
        if not ipn_id:
            return jsonify({"error": "Failed to register PesaPal IPN URL"}), 500
            
        # Prep user info
        user_profile = rtdb.reference(f'users/{user_id}').get() or {} if user_id != 'guest' else {}
        email = user_profile.get('email', session.get('user_email', 'customer@example.com'))
        phone = user_profile.get('phone', '0700000000')
        full_name = user_profile.get('full_name', 'Valued Farmer')
        
        names = full_name.split(' ', 1)
        first_name = names[0]
        last_name = names[1] if len(names) > 1 else 'Farmer'
        
        # Submit transaction
        order_res = submit_pesapal_order(
            token=token,
            order_reference=merchant_ref,
            amount=amount,
            currency=currency,
            description=f"Farmerman Systems: {plan_details['name']}",
            callback_url=callback_url,
            ipn_id=ipn_id,
            email=email,
            phone=phone,
            first_name=first_name,
            last_name=last_name
        )
        
        if order_res and isinstance(order_res, dict) and order_res.get('redirect_url'):
            return jsonify({'redirect_url': order_res['redirect_url']})
            
        return jsonify({"error": "Failed to initiate transaction with PesaPal", "details": order_res}), 500
    except Exception as e:
        print(f"PesaPal checkout creation failed: {e}")
        return jsonify({"error": "Server error during PesaPal initialization"}), 500

@app.route('/api/create-pesapal-session', methods=['POST'])
@token_required
def api_create_pesapal_session():
    """Allows the mobile app to initiate a PesaPal payment session."""
    try:
        data = request.json or {}
        plan_id = data.get('plan', 'bronze')
        category = data.get('category')
        
        plan_details = get_plan_price(plan_id, category)
        amount = float(plan_details['kes'])
        currency = "KES"
        
        uid = request.user['uid']
        # Generate merchant reference
        merchant_ref = f"PESA_{uid[-6:]}_{int(time.time())}"
        
        # Save pending transaction state
        rtdb.reference(f'pending_transactions/{merchant_ref}').set({
            'user_id': uid,
            'amount': amount,
            'plan_id': plan_id,
            'category': category or '',
            'status': 'awaiting_payment',
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'platform': 'mobile_app'
        })
        
        # Determine host url dynamically
        host_url = request.host_url.rstrip('/')
        callback_url = f"{host_url}/pesapal-callback"
        
        # Authenticate and obtain token
        token = get_pesapal_token()
        if not token:
            return jsonify({"error": "Failed to authenticate with PesaPal API"}), 500
            
        ipn_id = get_pesapal_ipn_id(token, host_url)
        if not ipn_id:
            return jsonify({"error": "Failed to register PesaPal IPN URL"}), 500
            
        # Prep user info
        user_profile = rtdb.reference(f'users/{uid}').get() or {}
        email = user_profile.get('email', 'customer@example.com')
        phone = user_profile.get('phone', '0700000000')
        full_name = user_profile.get('full_name', 'Valued Farmer')
        
        names = full_name.split(' ', 1)
        first_name = names[0]
        last_name = names[1] if len(names) > 1 else 'Farmer'
        
        # Submit transaction
        order_res = submit_pesapal_order(
            token=token,
            order_reference=merchant_ref,
            amount=amount,
            currency=currency,
            description=f"Farmerman Systems: {plan_details['name']}",
            callback_url=callback_url,
            ipn_id=ipn_id,
            email=email,
            phone=phone,
            first_name=first_name,
            last_name=last_name
        )
        
        if order_res and isinstance(order_res, dict) and order_res.get('redirect_url'):
            return jsonify({
                'redirect_url': order_res['redirect_url'],
                'merchant_reference': merchant_ref
            })
            
        return jsonify({"error": "Failed to initiate transaction with PesaPal", "details": order_res}), 500
    except Exception as e:
        print(f"PesaPal API session creation failed: {e}")
        return jsonify({"error": "Server error during PesaPal initialization"}), 500

@app.route('/pesapal-callback')
def pesapal_callback():
    tracking_id = request.args.get('OrderTrackingId')
    merchant_reference = request.args.get('OrderMerchantReference')
    
    if not tracking_id or not merchant_reference:
        flash("Invalid PesaPal callback details.", "warning")
        return redirect(url_for('pricing'))
        
    try:
        # Check transaction state
        token = get_pesapal_token()
        status_res = get_pesapal_transaction_status(token, tracking_id)
        
        payment_status = str(status_res.get('payment_status_description', '')).upper()
        
        if payment_status == 'COMPLETED':
            pending_txn = rtdb.reference(f'pending_transactions/{merchant_reference}').get()
            if pending_txn:
                uid = pending_txn.get('user_id')
                plan_id = pending_txn.get('plan_id')
                amount = float(status_res.get('amount', pending_txn.get('amount', 0)))
                category = pending_txn.get('category')
                
                record_successful_transaction(uid, plan_id, amount, "PesaPal", tracking_id)
                if plan_id == 'market_intel' and category and uid != 'guest':
                    try:
                        rtdb.reference(f'users/{uid}').update({'organization_category': category})
                    except Exception as ex:
                        print(f"Error saving category to database: {ex}")
                
                if plan_id == 'market_intel':
                    session['market_intel_guest'] = True
                        
                rtdb.reference(f'pending_transactions/{merchant_reference}').delete()
                
            flash(f"Payment successful! Welcome to your new plan.", "success")
            return redirect(url_for('payment_success'))
            
        elif payment_status in ['FAILED', 'INVALID']:
            flash(f"Payment failed: {status_res.get('description', 'Rejected by provider')}", "danger")
            return redirect(url_for('payment_failed'))
            
        # Processing state
        flash("Payment is currently processing. You will receive access as soon as the provider updates the state.", "info")
        return redirect(url_for('billing_history'))
    except Exception as e:
        print(f"PesaPal Callback verification error: {e}")
        flash("System error validating payment state. Please contact administrator.", "danger")
        return redirect(url_for('pricing'))

@app.route('/pesapal-ipn', methods=['GET', 'POST'])
def pesapal_ipn_listener():
    # Fetch parameters from both GET request query arguments and POST json body
    # PesaPal v3 uses OrderTrackingId, while some legacy or custom integrations use pesapal_transaction_tracking_id
    tracking_id = request.args.get('OrderTrackingId') or request.args.get('pesapal_transaction_tracking_id') or request.values.get('pesapal_transaction_tracking_id')
    merchant_reference = request.args.get('OrderMerchantReference') or request.args.get('pesapal_merchant_reference') or request.values.get('pesapal_merchant_reference')
    
    if request.is_json and not tracking_id:
        tracking_id = request.json.get('OrderTrackingId') or request.json.get('pesapal_transaction_tracking_id')
        merchant_reference = request.json.get('OrderMerchantReference') or request.json.get('pesapal_merchant_reference')
        
    print(f"PesaPal IPN Notification received. Tracking ID: {tracking_id}, Ref: {merchant_reference}")
    
    if tracking_id and merchant_reference:
        try:
            token = get_pesapal_token()
            status_res = get_pesapal_transaction_status(token, tracking_id)
            payment_status = str(status_res.get('payment_status_description', '')).upper()
            
            if payment_status == 'COMPLETED':
                pending_txn = rtdb.reference(f'pending_transactions/{merchant_reference}').get()
                if pending_txn:
                    uid = pending_txn.get('user_id')
                    plan_id = pending_txn.get('plan_id')
                    amount = float(status_res.get('amount', pending_txn.get('amount', 0)))
                    category = pending_txn.get('category')
                    
                    record_successful_transaction(uid, plan_id, amount, "PesaPal", tracking_id)
                    if plan_id == 'market_intel' and category and uid != 'guest':
                        try:
                            rtdb.reference(f'users/{uid}').update({'organization_category': category})
                        except Exception as ex:
                            print(f"Error saving category to database: {ex}")
                            
                    rtdb.reference(f'pending_transactions/{merchant_reference}').delete()
                    
            return jsonify({
                "orderNotificationType": "IPNCHANGE",
                "orderTrackingId": tracking_id,
                "orderMerchantReference": merchant_reference,
                "status": 200
            }), 200
        except Exception as e:
            print(f"PesaPal IPN processing failed: {e}")
            return jsonify({"error": "Failed to process notification"}), 500
            
    return jsonify({"error": "Invalid notification parameters"}), 400

# --- SUCCESS PAGE ---
@app.route('/success')
def payment_success(): 
    if session.get('pending_plan') == 'market_intel':
        session['market_intel_guest'] = True
    return render_template('payments/payment_success.html')

#==========================================
# STATIC PAGES & ERRORS
# ==========================================
@app.route('/')
def home(): 
    return render_template('home.html')

@app.route('/about')
def about_us(): 
    return render_template('about us.html')

@app.route('/impact')
def impact_initiatives(): 
    return render_template('impact&initiatives.html')

@app.route('/services')
def services(): 
    return render_template('our services.html')

@app.route('/privacy-policy')
def privacy_policy(): 
    return render_template('privacy policy.html')

@app.route('/manifest.json')
def serve_manifest():
    return send_from_directory('static', 'manifest.json')

@app.route('/sw.js')
def serve_sw():
    return send_from_directory('static', 'sw.js')


# DEDUPLICATED PRICING ROUTE
@app.route('/pricing')
def pricing():
    # Notice the "payments/" folder prefix here
    return render_template('payments/pricing_subscription.html')

@app.route('/terms-of-service')
def terms_of_service(): 
    return render_template('terms of service.html')

@app.route('/refund-policy')
def refund_policy(): 
    return render_template('subscription&refund policy.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact_us(): 
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        subject = request.form.get('subject') 
        message = request.form.get('message')
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        rtdb.reference('contact_inquiries').push({'name': name, 'email': email, 'subject': subject, 'message': message, 'timestamp': timestamp})
        
        current_year = datetime.now().year
        user_html = render_template('email_user_confirmation.html', name=name, message=message, year=current_year)
        admin_html = render_template('email_admin_notification.html', name=name, email=email, subject=subject, message=message, timestamp=timestamp)
        
        # BUG FIX: Use app.config['MAIL_USERNAME'] instead of undefined MAIL_USERNAME
        threading.Thread(target=send_async_emails, args=(email, app.config['MAIL_USERNAME'], user_html, admin_html, name, message, subject)).start()
        
        flash("Message sent! Check your email for a confirmation receipt.", "success")
        return redirect(url_for('contact_us'))
        
    return render_template('contact us.html')

@app.route('/delete_account', methods=['POST'])
@login_required
def delete_account():
    user_id = session.get('user_id')
    try:
        # 1. Remove from Firebase
        rtdb.reference(f'users/{user_id}').delete()
        
        # 2. Clear local session
        session.clear()
        
        flash("Your account has been permanently deleted.", "info")
        # Change 'index' to whatever your landing page function is called
        # The logs suggest 'login' might be your intended landing page
        return redirect(url_for('login')) 
        
    except Exception as e:
        print(f"Error deleting account: {e}")
        flash("An error occurred.", "danger")
        # The logs explicitly said use 'account_settings'
        return redirect(url_for('account_settings'))
    
@app.route('/diagnostics', methods=['GET'])
def diagnostics():
    health_data = {
        "status": "Healthy",
        "system": "Farmerman Systems",
        "timestamp": datetime.now().strftime("%Y-%b-%d %H:%M:%S"),
        "version": "2.1.0",
        "uptime": "99.9%" 
    }
    return render_template('diagnostics.html', data=health_data)
# ==========================================
# REAL-TIME CHAT SYSTEM (Strict Privacy & Online-Only)
# ==========================================
@app.route('/chat/dashboard')
@login_required
@premium_required
# @premium_required # Uncomment if you want this protected
def chat_dashboard():
    """Chat dashboard showing all users and their current online status."""
    current_uid = session.get('user_id')
    all_users = rtdb.reference('users').get() or {}
    
    online_contacts = []
    for uid, data in all_users.items():
        # WE NO LONGER FILTER OUT OFFLINE USERS HERE!
        if uid != current_uid:
            online_contacts.append({
                'uid': uid,
                'name': data.get('full_name', 'Farmer'),
                'role': data.get('role', 'buyer'),
                # We pass the status to the frontend instead of filtering
                'is_online': uid in online_users 
            })
            
    # Pointing exactly to the file inside the chat folder
    return render_template('chat/dashboard.html', online_contacts=online_contacts)

@app.route('/chat')
@login_required
def chat_home():
    """Renders the main chat UI with smart sorting and unread badges."""
    current_uid = session.get('user_id')
    
    auto_open_uid = request.args.get('target_uid')
    auto_open_name = request.args.get('target_name')
    
    all_users = rtdb.reference('users').get() or {}
    all_chats = rtdb.reference('chats').get() or {}
    
    # FETCH UNREAD COUNTS FOR THE CURRENT USER
    my_unread_counts = rtdb.reference(f'unread_counts/{current_uid}').get() or {}
    
    current_user_profile = all_users.get(current_uid, {})
    my_location = current_user_profile.get('location', '').strip().lower()
    
    contacts = []
    for uid, data in all_users.items():
        if uid != current_uid:
            room_id = f"room_{min(str(current_uid), str(uid))}_{max(str(current_uid), str(uid))}"
            has_talked_before = 1 if room_id in all_chats else 0
            their_location = data.get('location', '').strip().lower()
            is_same_location = 1 if (my_location and their_location == my_location) else 0
            
            # Extract the specific unread count for this contact
            unread_count = my_unread_counts.get(uid, 0)

            contacts.append({
                'uid': uid,
                'name': data.get('full_name', 'Farmer'),
                'role': data.get('role', 'client'),
                'is_online': uid in online_users,
                'has_talked': has_talked_before,
                'same_location': is_same_location,
                'unread': unread_count # <-- NEW DATA POINT
            })
            
    # Sort logic: Unread messages float to the very top, then Talked, then Location, then Online
    contacts.sort(key=lambda x: (x['unread'] > 0, x['has_talked'], x['same_location'], x['is_online']), reverse=True)
            
    return render_template(
        'chat/index.html', 
        contacts=contacts, 
        current_uid=current_uid,
        auto_open_uid=auto_open_uid,   
        auto_open_name=auto_open_name  
    )
    
@app.route('/api/chat/upload', methods=['POST'])
@login_required
@premium_required
def upload_chat_media():
    """Handles image, video, and audio uploads to Cloud Storage."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Empty file'}), 400
        
    file_url = upload_to_firebase_storage(file, 'chat_media')
    
    if file_url:
        return jsonify({'url': file_url, 'type': file.content_type}), 200
    return jsonify({'error': 'Cloud upload failed'}), 500


# --- SOCKET.IO EVENTS ---

@socketio.on('connect')
def handle_connect():
    uid = session.get('user_id')
    if uid:
        online_users[uid] = request.sid
        
        # NEW: Join a personal room using the user's ID. 
        # This allows the server to send global notifications (like unread badges) 
        # directly to this user, regardless of what chat they are looking at.
        join_room(uid) 
        
        # 1. Update the green dots for everyone
        emit('user_status', {'uid': uid, 'status': 'online'}, broadcast=True)
        # 2. Tell everyone to refresh their sidebar
        emit('refresh_contacts', broadcast=True, include_self=False)

@socketio.on('disconnect')
def handle_disconnect():
    uid = session.get('user_id')
    if uid and uid in online_users:
        del online_users[uid]
        # 1. Turn the dot grey for everyone
        emit('user_status', {'uid': uid, 'status': 'offline'}, broadcast=True)
        # 2. Tell everyone to refresh sidebar
        emit('refresh_contacts', broadcast=True)

@socketio.on('join_chat')
def handle_join_chat(data):
    uid1 = session.get('user_id')
    uid2 = data.get('target_uid')
    if not uid2: return

    room = f"room_{min(str(uid1), str(uid2))}_{max(str(uid1), str(uid2))}"
    join_room(room)
    
    history = rtdb.reference(f'chats/{room}').get() or {}
    messages = []
    
    # SMART FETCH: Only send messages that the user hasn't deleted for themselves
    for msg in history.values():
        if uid1 not in msg.get('deleted_for', []):
            messages.append(msg)
            
    # NEW: RESET UNREAD COUNTER
    # Since the user just opened this chat, set their unread count from this sender to 0
    rtdb.reference(f'unread_counts/{uid1}/{uid2}').set(0)
            
    emit('chat_history', messages)

@socketio.on('clear_chat')
def handle_clear_chat(data):
    uid = session.get('user_id')
    target_uid = data.get('target_uid')
    mode = data.get('mode', 'me') # Can be 'me' or 'all'
    if not target_uid: return
    
    room = f"room_{min(str(uid), str(target_uid))}_{max(str(uid), str(target_uid))}"
    
    if mode == 'all':
        # Wipe it completely from the database
        rtdb.reference(f'chats/{room}').delete()
        emit('chat_cleared', {'mode': 'all'}, room=room)
    else:
        # CLEAR FOR ME ONLY
        messages_ref = rtdb.reference(f'chats/{room}')
        messages = messages_ref.get() or {}
        
        # Loop through existing messages and tag them as hidden for this specific user
        for msg_key, msg_data in messages.items():
            deleted_for = msg_data.get('deleted_for', [])
            if uid not in deleted_for:
                deleted_for.append(uid)
                messages_ref.child(msg_key).update({'deleted_for': deleted_for})
                
        # Emit the clear screen event ONLY to the person who clicked the button
        emit('chat_cleared', {'mode': 'me'}, to=request.sid)
        
@socketio.on('send_message')
def handle_send_message(data):
    sender_id = session.get('user_id')
    receiver_id = data.get('receiver_id')
    if not receiver_id: return

    room = f"room_{min(str(sender_id), str(receiver_id))}_{max(str(sender_id), str(receiver_id))}"
    
    # NEW: Define East Africa Time (UTC+3)
    eat_timezone = timezone(timedelta(hours=3))
    
    message_data = {
        'sender_id': sender_id,
        'text': data.get('text', ''),
        'media_url': data.get('media_url'),
        'media_type': data.get('media_type'),
        # NEW: Apply the timezone to the current time
        'timestamp': datetime.now(eat_timezone).strftime("%H:%M")
    }
    
    # 1. Save Message to Database
    rtdb.reference(f'chats/{room}').push(message_data)
    
    # 2. INCREMENT UNREAD COUNTER
    unread_ref = rtdb.reference(f'unread_counts/{receiver_id}/{sender_id}')
    current_unread = unread_ref.get() or 0
    new_unread_count = current_unread + 1
    unread_ref.set(new_unread_count)
    
    # 3. Broadcast the message to the active chat room
    emit('receive_message', message_data, room=room)
    
    # 4. Global UI Notification
    emit('update_unread_badge', {
        'sender_id': sender_id, 
        'count': new_unread_count
    }, room=receiver_id)
    
@socketio.on('typing')
def handle_typing(data):
    receiver_id = data.get('receiver_id')
    sender_id = session.get('user_id')
    room = f"room_{min(str(sender_id), str(receiver_id))}_{max(str(sender_id), str(receiver_id))}"
    emit('display_typing', {'sender_id': sender_id}, room=room, include_self=False)

@socketio.on('stop_typing')
def handle_stop_typing(data):
    receiver_id = data.get('receiver_id')
    sender_id = session.get('user_id')
    room = f"room_{min(str(sender_id), str(receiver_id))}_{max(str(sender_id), str(receiver_id))}"
    emit('hide_typing', {'sender_id': sender_id}, room=room, include_self=False)
# ==========================================
# ENTERPRISE EXPANSION HUBS (Live Database Integration)
# ==========================================

@app.route('/deal-room')
@login_required
@premium_required
def deal_room():
    """Exclusive portal for Gold/Investor tier to view bankable projects."""
    # Fetch live deals from Firebase
    deals_data = rtdb.reference('deals').get() or {}
    deals = [{'id': k, **v} for k, v in deals_data.items()]
    
    return render_template('deal_room.html', deals=deals)

@app.route('/api-docs')
@premium_required
def api_docs():
    """Developer documentation for Enterprise clients."""
    return render_template('api_docs.html')
# In main.py
@app.route('/climate')
@login_required
@premium_required
def climate_hub():
    """Weather and Climate Smart Agriculture dashboard."""
    user_id = session.get('user_id')
    
    # Fetch ONLY this user's personal alerts from Firebase
    climate_data = rtdb.reference(f'climate_alerts/{user_id}').get()
    
    # Format the data depending on how Firebase returned it
    alerts = []
    if isinstance(climate_data, dict):
        alerts = [{'id': k, **v} for k, v in climate_data.items()]
    elif isinstance(climate_data, list):
        # Firebase sometimes stores .set() lists as actual arrays
        alerts = [a for a in climate_data if a is not None]
        
    # Sort newest first
    alerts.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
    
    return render_template('climate.html', alerts=alerts)

# Make sure you have these imports at the top of your main.py!
from flask import request, jsonify
from logic import analyze_weather_and_generate_alerts, update_firebase_alerts
# In main.py
@app.route('/api/climate/analyze', methods=['POST'])
@token_required
def analyze_climate():
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"status": "error", "message": "Invalid JSON"}), 400
    
    try:
        # Get the currently logged-in user's ID
        user_id = session.get('user_id')
        
        temp = float(data.get('temp', 0))
        humidity = float(data.get('humidity', 0))
        wind = float(data.get('wind', 0))
        condition = str(data.get('condition', ''))
        region = str(data.get('region', 'Unknown Region'))
        
        # Generate the alerts
        new_alerts = analyze_weather_and_generate_alerts(temp, humidity, wind, condition, region)
        
        # Save them ONLY to this user's profile!
        update_firebase_alerts(user_id, new_alerts)
        
        return jsonify({"status": "success", "message": "Updated personal alerts!"}), 200
        
    except Exception as e:
        print(f"Engine Error: {e}")
        return jsonify({"status": "error", "message": "Failed to analyze"}), 500
    
@app.route('/admin/add-climate-alert', methods=['POST'])
@admin_required
def add_climate_alert():
    """Admin route to publish new agronomic weather alerts."""
    try:
        rtdb.reference('climate_alerts').push({
            'region': request.form.get('region'),
            'temp': request.form.get('temp'),
            'condition': request.form.get('condition'),
            'humidity': request.form.get('humidity'),
            'wind': request.form.get('wind'),
            'alert_type': request.form.get('alert_type'), # 'warning' or 'success'
            'title': request.form.get('title'),
            'advice': request.form.get('advice'),
            'timestamp': datetime.now().strftime("%B %d, %Y - %H:%M")
        })
        flash("Climate alert published to the Hub!", "success")
    except Exception as e:
        flash(f"Error publishing climate alert: {e}", "danger")
        
    return redirect(url_for('admin_dashboard'))


@app.route('/insights')
@premium_required
def insights():
    """SEO-friendly thought leadership and market analysis blog."""
    # Fetch live articles from Firebase
    insights_data = rtdb.reference('insights').get() or {}
    articles = [{'id': k, **v} for k, v in insights_data.items()]
    
    return render_template('insights.html', articles=articles) 


@app.route('/insights/<article_id>')
def read_insight(article_id):
    """Fetches and displays a single full-length article."""
    # Fetch the specific article from Firebase
    article = rtdb.reference(f'insights/{article_id}').get()
    
    if not article:
        flash("Article not found.", "warning")
        return redirect(url_for('insights'))
        
    # Pass the data to the reader template
    return render_template('read_article.html', article=article)

# ==========================================
# ADMIN POSTING ROUTES FOR ENTERPRISE HUBS
# ==========================================
@app.route('/admin/add-deal', methods=['POST'])
@admin_required
def add_deal():
    """Admin route to publish new investment opportunities to the Deal Room."""
    try:
        rtdb.reference('deals').push({
            'title': request.form.get('title'),
            'ask': request.form.get('ask'),       # e.g., "$50,000"
            'roi': request.form.get('roi'),       # e.g., "14%"
            'sector': request.form.get('sector'), # e.g., "Agri-Tech"
            'risk': request.form.get('risk'),     # e.g., "Low"
            'status': request.form.get('status', 'Reviewing') # e.g., "Funding"
        })
        flash("New investment deal published to the Deal Room!", "success")
    except Exception as e:
        flash(f"Error posting deal: {e}", "danger")
        
    return redirect(url_for('admin_dashboard'))
@app.route('/admin/add-insight', methods=['POST'])
@admin_required
def add_insight():
    """Admin route to publish new Thought Leadership articles with Multimedia."""
    try:
        # 1. Grab files from the form
        image_file = request.files.get('image_file')
        video_file = request.files.get('video_file')
        audio_file = request.files.get('audio_file')
        
        # 2. Upload to Firebase Storage (if the file was provided)
        img_url = upload_to_firebase_storage(image_file, 'insights_media') if image_file and image_file.filename else "https://images.unsplash.com/photo-1592982537447-6f296cb31454?auto=format&fit=crop&w=600&q=80" # Fallback image
        video_url = upload_to_firebase_storage(video_file, 'insights_media') if video_file and video_file.filename else None
        audio_url = upload_to_firebase_storage(audio_file, 'insights_media') if audio_file and audio_file.filename else None

        # 3. Save to Realtime Database
        rtdb.reference('insights').push({
            'title': request.form.get('title'),
            'date': datetime.now().strftime("%B %d, %Y"),
            'category': request.form.get('category'),
            'read_time': request.form.get('read_time'), 
            'img': img_url,
            'video_url': video_url,
            'audio_url': audio_url,
            'summary': request.form.get('summary'),
            'content': request.form.get('content')
        })
        flash("Thought Leadership article with media published successfully!", "success")
    except Exception as e:
        flash(f"Error publishing insight: {e}", "danger")
        
    return redirect(url_for('admin_dashboard'))
# ==========================================
# DIGITAL SAVINGS & TABLE BANKING SYSTEM (MULTI-GROUP)
# ==========================================

@app.route('/banking')
@login_required
def banking_dashboard():
    uid = session.get('user_id')
    eat_tz = timezone(timedelta(hours=3))
    current_time = datetime.now(eat_tz)
    
    # 1. Fetch User's Banking Account
    account_ref = rtdb.reference(f'banking_accounts/{uid}')
    account = account_ref.get()
    
    # Initialize if totally new
    if not account:
        account = {'standard_savings': {}, 'emergency_fund': 0.0, 'groups': {}}
        account_ref.set(account)
        
    # Ensure legacy accounts are converted to the multi-group schema
    if 'group_id' in account:
        old_group = account.pop('group_id')
        old_savings = account.pop('standard_savings', 0.0)
        account['groups'] = {old_group: True} if old_group else {}
        account['standard_savings'] = {old_group: old_savings} if old_group else {}
        account_ref.set(account)

    # 2. Fetch All Groups & Categorize Them
    all_groups = rtdb.reference('banking_groups').get() or {}
    my_groups = []
    available_groups = []
    
    for gid, gdata in all_groups.items():
        if not isinstance(gdata, dict): continue
        
        gdata['id'] = gid
        gdata['member_count'] = len(gdata.get('members', {}))
        
        # Calculate maturity
        try:
            maturity_date = datetime.strptime(gdata['cycle_end_date'], "%Y-%m-%d").replace(tzinfo=eat_tz)
            time_diff = maturity_date - current_time
            gdata['days_to_maturity'] = max(0, time_diff.days)
            gdata['can_withdraw'] = gdata['days_to_maturity'] == 0
        except Exception:
            gdata['days_to_maturity'] = 0
            gdata['can_withdraw'] = False
            
        # Is the user in this group?
        if uid in gdata.get('members', {}):
            # Fetch their specific balance for this group
            gdata['my_balance'] = account.get('standard_savings', {}).get(gid, 0.0)
            my_groups.append(gdata)
        else:
            available_groups.append(gdata)

    # 3. Fetch Transactions
    transactions_data = rtdb.reference(f'banking_transactions/{uid}').get()
    tx_list = []
    if isinstance(transactions_data, dict):
        tx_list = [v for v in transactions_data.values() if isinstance(v, dict)]
    elif isinstance(transactions_data, list):
        tx_list = [v for v in transactions_data if isinstance(v, dict)]
        
    tx_list.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
    tx_list = tx_list[:10]

    # 4. Fetch Loans
    loans_data = rtdb.reference(f'banking_loans/{uid}').get()
    loan_list = []
    if isinstance(loans_data, dict):
        loan_list = [v for v in loans_data.values() if isinstance(v, dict)]
    elif isinstance(loans_data, list):
        loan_list = [v for v in loans_data if isinstance(v, dict)]

    return render_template(
        'banking/dashboard.html',
        account=account,
        my_groups=my_groups,
        available_groups=available_groups,
        transactions=tx_list,
        loans=loan_list
    )

@app.route('/banking/create-group', methods=['POST'])
@login_required
def banking_create_group():
    uid = session.get('user_id')
    eat_tz = timezone(timedelta(hours=3))
    
    group_name = request.form.get('group_name')
    duration_months = int(request.form.get('duration_months', 6))
    maturity_date = datetime.now(eat_tz) + timedelta(days=duration_months * 30)
    
    import uuid
    group_id = f"group_{uuid.uuid4().hex[:8]}"
    
    rtdb.reference(f'banking_groups/{group_id}').set({
        'name': group_name,
        'description': request.form.get('description'),
        'cycle_end_date': maturity_date.strftime("%Y-%m-%d"),
        'creator_id': uid,
        'members': {uid: True}
    })
    
    # Add to user's group list and initialize a 0.0 balance for it
    rtdb.reference(f'banking_accounts/{uid}/groups/{group_id}').set(True)
    rtdb.reference(f'banking_accounts/{uid}/standard_savings/{group_id}').set(0.0)
    
    flash(f"Successfully created '{group_name}'.", "success")
    return redirect(url_for('banking_dashboard'))

@app.route('/banking/join-group/<group_id>', methods=['POST'])
@login_required
def banking_join_group(group_id):
    uid = session.get('user_id')
    rtdb.reference(f'banking_groups/{group_id}/members/{uid}').set(True)
    rtdb.reference(f'banking_accounts/{uid}/groups/{group_id}').set(True)
    rtdb.reference(f'banking_accounts/{uid}/standard_savings/{group_id}').set(0.0)
    flash("Successfully joined the savings group!", "success")
    return redirect(url_for('banking_dashboard'))

# ---------------------------------------------------------
# BANKING DEPOSITS (M-PESA)
# ---------------------------------------------------------
@app.route('/banking/process-deposit', methods=['POST'])
@login_required
def banking_process_deposit():
    """Handles M-Pesa STK Push for Banking Deposits."""
    uid = session.get('user_id')
    eat_tz = timezone(timedelta(hours=3))
    
    # The fund type can be 'emergency_fund' OR 'group_XYZ123'
    fund_type_or_group = request.form.get('fund_type')
    phone_number = request.form.get('phone_number')
    
    try:
        amount = float(request.form.get('amount', 0))
    except ValueError:
        flash("Invalid amount.", "danger")
        return redirect(url_for('banking_dashboard'))

    if amount <= 0:
        flash("Amount must be greater than zero.", "warning")
        return redirect(url_for('banking_dashboard'))

    try:
        res = initiate_stk_push(phone_number, int(amount))
        
        if res and res.get('ResponseCode') == '0':
            checkout_id = res.get("CheckoutRequestID")
            rtdb.reference(f'pending_transactions/{checkout_id}').set({
                'user_id': uid, 
                'amount': amount, 
                'fund_type': fund_type_or_group,
                'status': 'awaiting_payment',
                'tx_type': 'banking_deposit', 
                'timestamp': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S")
            })
            
            # CRITICAL FIX: We must pass source='banking' here!
            return redirect(url_for('payment_processing', checkout_id=checkout_id, source='banking'))
            
        flash(f"M-Pesa Error: {res.get('errorMessage', 'Service unavailable.')}", "danger")
    except Exception as e:
        flash("M-Pesa Gateway is currently unstable. Please try again.", "danger")
        
    return redirect(url_for('banking_dashboard'))

# ---------------------------------------------------------
# WITHDRAWALS & LOANS
# ---------------------------------------------------------

@app.route('/banking/withdraw', methods=['POST'])
@login_required
def banking_process_withdraw():
    """Deducts funds and securely processes withdrawal requests."""
    uid = session.get('user_id')
    eat_tz = timezone(timedelta(hours=3))
    
    target_account = request.form.get('fund_type') # 'emergency_fund' or 'group_XYZ123'
    phone_number = request.form.get('phone_number')
    
    try:
        amount = float(request.form.get('amount', 0))
    except ValueError:
        flash("Invalid amount.", "danger")
        return redirect(url_for('banking_dashboard'))
    
    account_ref = rtdb.reference(f'banking_accounts/{uid}')
    account = account_ref.get()
    
    # Route logic based on where they are withdrawing from
    if target_account == 'emergency_fund':
        current_balance = account.get('emergency_fund', 0.0)
        if amount > current_balance:
            flash("Insufficient emergency funds.", "danger")
            return redirect(url_for('banking_dashboard'))
        account_ref.update({'emergency_fund': current_balance - amount})
        log_name = "Emergency Fund"
        
    else:
        # It's a standard savings group
        group_id = target_account
        current_balance = account.get('standard_savings', {}).get(group_id, 0.0)
        
        if amount > current_balance:
            flash("Insufficient group savings.", "danger")
            return redirect(url_for('banking_dashboard'))
            
        # Enforce Lock-in for this specific group
        group = rtdb.reference(f'banking_groups/{group_id}').get()
        maturity_date = datetime.strptime(group['cycle_end_date'], "%Y-%m-%d").replace(tzinfo=eat_tz)
        if datetime.now(eat_tz) < maturity_date:
            flash(f"Savings in '{group['name']}' are locked until {group['cycle_end_date']}.", "danger")
            return redirect(url_for('banking_dashboard'))

        # Deduct balance
        new_balance = current_balance - amount
        rtdb.reference(f'banking_accounts/{uid}/standard_savings/{group_id}').set(new_balance)
        log_name = group['name']
    
    # Log withdrawal
    rtdb.reference(f'banking_transactions/{uid}').push({
        'type': 'withdraw',
        'fund_type': log_name,
        'amount': amount,
        'destination': phone_number,
        'timestamp': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S")
    })
    
    flash(f"Withdrawal of KES {amount:,.2f} is being processed to {phone_number}.", "success")
    return redirect(url_for('banking_dashboard'))
@app.route('/banking/loan', methods=['POST'])
@login_required
def banking_request_loan():
    uid = session.get('user_id')
    eat_tz = timezone(timedelta(hours=3))
    
    loan_data = {
        'amount': float(request.form.get('amount', 0)),
        'type': request.form.get('loan_type'),
        'reason': request.form.get('reason'),
        'provider': request.form.get('provider'),
        'phone_number': request.form.get('phone_number'), # <-- NEW: Captures the verification number
        'status': 'Pending Review' if request.form.get('provider') == 'system' else 'Referred',
        'requested_at': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S")
    }
    rtdb.reference(f'banking_loans/{uid}').push(loan_data)
    
    if loan_data['provider'] == 'external_bank':
        flash("Request logged. Visit your nearest partner bank branch with your ID.", "info")
    else:
        # Inform them that they will receive a call
        flash("System loan request submitted. An admin will call the provided number shortly to verify.", "success")
        
    return redirect(url_for('banking_dashboard'))


# ==========================================
# FLUTTER MOBILE APP APIs (JSON Endpoints)
# ==========================================

@app.route('/api/banking/dashboard', methods=['GET'])
@token_required # Validates the mobile app's Firebase Bearer token
def api_banking_dashboard():
    """Returns the user's banking data as JSON for the Flutter app."""
    try:
        uid = request.user['uid']
        
        # 1. Fetch Account Balances
        account_ref = rtdb.reference(f'banking_accounts/{uid}')
        account = account_ref.get() or {'emergency_fund': 0.0, 'standard_savings': {}}
        
        emergency_fund = float(account.get('emergency_fund', 0.0))
        
        # Calculate sum of all savings groups
        standard_savings = account.get('standard_savings', {})
        total_savings = sum(float(v) for v in standard_savings.values())
        
        # 2. Fetch Transactions
        txns_data = rtdb.reference(f'banking_transactions/{uid}').get() or {}
        tx_list = []
        if isinstance(txns_data, dict):
            tx_list = [v for v in txns_data.values() if isinstance(v, dict)]
            
        # Sort newest first
        tx_list.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        
        return jsonify({
            "status": "success",
            "emergency_fund": emergency_fund,
            "total_savings": total_savings,
            "recent_transactions": tx_list[:10] # Return the 10 most recent
        }), 200
        
    except Exception as e:
        print(f"API Banking Error: {e}")
        return jsonify({"error": "Failed to load banking data"}), 500


@app.route('/api/process-mpesa', methods=['POST'])
@token_required 
def api_process_mpesa():
    """Triggers Safaricom STK push from the Flutter app."""
    try:
        data = request.json
        phone = data.get('phone_number')
        amount = int(float(data.get('amount', 0)))
        fund_type = data.get('fund_type', 'emergency_fund')
        tx_type = data.get('tx_type', 'banking_deposit') # Can be 'banking_deposit' or 'subscription'
        plan_id = data.get('plan_id', 'pro')
        
        uid = request.user['uid']
        eat_tz = timezone(timedelta(hours=3))

        if amount <= 0:
            return jsonify({"error": "Amount must be greater than zero."}), 400

        # Trigger your actual mpesa.py logic
        res = initiate_stk_push(phone, amount)
        
        if res and res.get('ResponseCode') == '0':
            checkout_id = res.get("CheckoutRequestID")
            
            # Store in pending transactions so your mpesa-callback can match it
            rtdb.reference(f'pending_transactions/{checkout_id}').set({
                'user_id': uid, 
                'amount': amount, 
                'fund_type': fund_type,
                'plan_id': plan_id,
                'status': 'awaiting_payment',
                'tx_type': tx_type, 
                'timestamp': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S")
            })
            
            return jsonify({"status": "success", "checkout_id": checkout_id}), 200
            
        return jsonify({"error": res.get('errorMessage', 'Safaricom service unavailable.')}), 400
        
    except Exception as e:
        print(f"API STK Error: {e}")
        return jsonify({"error": "Gateway unstable. Please try again."}), 500


# ==========================================
# PREMIUM INSIGHTS API
# ==========================================
@app.route('/api/insights', methods=['GET'])
@token_required
def api_insights():
    """Fetches live thought leadership articles for the mobile app."""
    try:
        # Fetch live articles from Firebase
        insights_data = rtdb.reference('insights').get() or {}
        
        articles = []
        for key, val in insights_data.items():
            if isinstance(val, dict):
                # Attach the firebase key as the ID
                articles.append({'id': key, **val})
        
        # Optional: Sort articles by date (newest first) assuming you have a 'date' field
        articles.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        return jsonify({
            "status": "success", 
            "articles": articles
        }), 200
        
    except Exception as e:
        print(f"API Insights Error: {e}")
        return jsonify({"error": "Failed to load premium insights."}), 500


# ==========================================
# EXCLUSIVE DEAL ROOM API
# ==========================================
@app.route('/api/deal-room', methods=['GET'])
@token_required
def api_deal_room():
    """Exclusive portal for Gold/Investor tier to view bankable projects via mobile."""
    try:
        uid = request.user['uid']
        
        # 1. Security Check: Fetch user data to verify their tier
        user_data = rtdb.reference(f'users/{uid}').get() or {}
        user_tier = str(user_data.get('subscription_tier', 'free')).lower()
        user_role = str(user_data.get('role', 'client')).lower()
        
        # 2. Enforce the Paywall on the backend
        if user_tier not in ['gold', 'investor'] and user_role != 'admin':
            return jsonify({
                "status": "error", 
                "error": "Access Denied. Please upgrade to the Gold or Investor tier."
            }), 403

        # 3. Fetch live deals from Firebase
        deals_data = rtdb.reference('deals').get() or {}
        
        deals = []
        for key, val in deals_data.items():
            if isinstance(val, dict):
                deals.append({'id': key, **val})
                
        return jsonify({
            "status": "success", 
            "deals": deals
        }), 200
        
    except Exception as e:
        print(f"API Deal Room Error: {e}")
        return jsonify({"error": "Failed to load deal room data."}), 500
    
# ==========================================
# MOBILE APP LOAN REQUEST API
# ==========================================
@app.route('/api/banking/loan', methods=['POST'])
@token_required 
def api_request_loan():
    """Handles loan requests coming from the Flutter mobile app."""
    try:
        uid = request.user['uid']
        data = request.get_json(silent=True) or {}
        
        amount = float(data.get('amount', 0))
        reason = str(data.get('reason', 'General Agriculture'))
        
        if amount < 500:
            return jsonify({"error": "Minimum loan amount is KES 500"}), 400
            
        eat_tz = timezone(timedelta(hours=3))
        
        # 1. Fetch user info to attach to the loan request
        user_data = rtdb.reference(f'users/{uid}').get() or {}
        full_name = user_data.get('full_name', 'Unknown User')
        phone = user_data.get('phone', 'No Phone')

        # 2. Push the new loan request to Firebase
        loan_ref = rtdb.reference('loan_requests').push()
        loan_ref.set({
            'user_id': uid,
            'full_name': full_name,
            'phone': phone,
            'amount': amount,
            'reason': reason,
            'status': 'pending',  # Awaiting Admin Approval
            'timestamp': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S"),
            'platform': 'mobile_app'
        })
        
        return jsonify({
            "status": "success", 
            "message": "Loan requested successfully"
        }), 200
        
    except Exception as e:
        print(f"API Loan Request Error: {e}")
        return jsonify({"error": "Failed to process loan request"}), 500
    
    
    
   
   
   
    
# ==========================================
# MOBILE APP WITHDRAWAL REQUEST API
# ==========================================






@app.route('/api/banking/withdraw', methods=['POST'])
@token_required 
def api_request_withdrawal():
    """Handles withdrawal requests coming from the Flutter mobile app."""
    try:
        uid = request.user['uid']
        data = request.get_json(silent=True) or {}
        
        amount = float(data.get('amount', 0))
        phone = str(data.get('phone', ''))
        
        if amount <= 0:
            return jsonify({"error": "Invalid amount requested."}), 400
        if not phone:
            return jsonify({"error": "A receiving M-Pesa number is required."}), 400
            
        eat_tz = timezone(timedelta(hours=3))
        
        # 1. Fetch user info to attach to the request
        user_data = rtdb.reference(f'users/{uid}').get() or {}
        full_name = user_data.get('full_name', 'Unknown User')

        # 2. Add validation (Optional: Check if they actually have this money)
        # For a full banking app, you would check `user_data.get('total_savings') >= amount` here.

        # 3. Push the withdrawal request to Firebase
        withdraw_ref = rtdb.reference('withdrawal_requests').push()
        withdraw_ref.set({
            'user_id': uid,
            'full_name': full_name,
            'phone': phone,
            'amount': amount,
            'status': 'pending',  # Awaiting Admin to disburse funds
            'timestamp': datetime.now(eat_tz).strftime("%Y-%m-%d %H:%M:%S"),
            'platform': 'mobile_app'
        })
        
        return jsonify({
            "status": "success", 
            "message": "Withdrawal requested successfully"
        }), 200
        
    except Exception as e:
        print(f"API Withdrawal Error: {e}")
        return jsonify({"error": "Failed to process withdrawal"}), 500
        
        
        
        
        
        
        
        
# ==========================================
# ZIM BOT (WHATSAPP AGRICULTURAL ASSISTANT)
# ==========================================


@app.route("/webhook/twilio/zim-bot", methods=['POST'])
def zim_bot_webhook():
    incoming_msg = request.values.get('Body', '').strip()
    raw_phone = request.values.get('From', '') # e.g. "whatsapp:+263771234567"
    
    # Clean the phone number to match database format
    user_phone = raw_phone.replace('whatsapp:', '').replace('+', '').strip()

    # --- PERSONA & LOGIC DEFINITIONS ---
    GREETINGS = ['hi', 'hello', 'mhoroi', 'sanibonani', 'hey', 'start', 'tora', 'nzou', 'menu']
    persona_intro = "🇿🇼 *ZimBot: Professional Agricultural Advisor*\n\n"
    
    # Package Descriptions
    PRICING_MENU = (
        "Professional Packages:\n\n"
        "1️⃣ *Seed Package ($3/mo)*\n"
        "• Weekly Market Intelligence (Friday Summary)\n"
        "• Basic Price Tracking\n\n"
        "2️⃣ *Growth Package ($5/mo)*\n"
        "• Bi-weekly Intelligence (Monday & Friday)\n"
        "• Trend Analysis & Priority Alerts\n\n"
        "3️⃣ *Harvest Package ($10/mo)*\n"
        "• Full Intelligence (Monday, Wednesday, Friday)\n"
        "• Strategic Action Plans & Deep Market Insights\n"
        "• Capitalize on market shifts to maximize ROI\n\n"
        "Reply with *1*, *2*, or *3* to select your package and begin."
    )

    try:
        # 1. AUTHENTICATION: Identify User
        user_data = None
        all_users = rtdb.reference('users').get() or {}
        for uid, data in all_users.items():
            db_phone = str(data.get('phone', '')).replace('+', '').strip()
            if db_phone and (db_phone in user_phone or user_phone in db_phone):
                user_data = data
                break

        # 2. REGISTRATION FLOW: Handle Unknown Users
        if not user_data:
            pending_ref = rtdb.reference(f'pending_registrations/{user_phone}')
            pending = pending_ref.get()

            if not pending:
                if any(greet in incoming_msg.lower() for greet in GREETINGS):
                    reply_text = f"{persona_intro}Welcome to Zimbabwe's premier market intelligence network. To access our strategic data, please provide your *Full Name*."
                    pending_ref.set({'step': 'ask_name'})
                else:
                    reply_text = f"{persona_intro}Welcome! Please say 'Hi' or 'Start' to begin your registration and unlock live market intelligence."
            
            elif pending.get('step') == 'ask_name':
                full_name = incoming_msg.title()
                pending_ref.update({'full_name': full_name, 'step': 'ask_email'})
                reply_text = f"Thank you, {full_name}. What is your *Email Address* for receiving professional market reports?"
            
            elif pending.get('step') == 'ask_email':
                email = incoming_msg.strip().lower()
                if '@' not in email or '.' not in email:
                    reply_text = "Please provide a valid email address (e.g., advisor@gmail.com)."
                else:
                    full_name = pending.get('full_name')
                    new_uid = f"wa_{user_phone}"
                    rtdb.reference(f'users/{new_uid}').set({
                        'uid': new_uid,
                        'full_name': full_name,
                        'email': email,
                        'phone': user_phone,
                        'role': 'buyer',
                        'subscription_tier': 'free',
                        'subscription_status': 'inactive',
                        'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    })
                    pending_ref.delete()
                    # MANDATORY NEXT STEP: SUBSCRIPTION
                    rtdb.reference(f'pending_payments/{user_phone}').set({'step': 'select_package'})
                    reply_text = f"Welcome aboard, {full_name}. Profile created.\n\nTo access our live intelligence network, please {PRICING_MENU}"
            
            resp = MessagingResponse(); resp.message(reply_text); return str(resp)

        # 3. SUBSCRIPTION GATE: Registered but Inactive
        user_status = user_data.get('subscription_status', 'inactive')
        user_tier = user_data.get('subscription_tier', 'free')
        
        pay_ref = rtdb.reference(f'pending_payments/{user_phone}')
        pay_state = pay_ref.get()

        if user_status != 'active':
            # Force user into payment flow
            if not pay_state:
                pay_ref.set({'step': 'select_package'})
                reply_text = f"{persona_intro}Your account is currently inactive. To access market intelligence, please choose a subscription plan:\n\n{PRICING_MENU}"
            else:
                step = pay_state.get('step')
                if step == 'select_package':
                    selection = incoming_msg.strip()
                    plans = {'1': 'seed', '2': 'growth', '3': 'harvest'}
                    if selection in plans:
                        plan_id = plans[selection]
                        plan_name = plan_id.capitalize()
                        price_kes = SYSTEM_PRICING[plan_id]['kes']
                        pay_ref.update({'step': 'confirm_phone', 'plan_id': plan_id, 'amount': price_kes})
                        reply_text = f"Selected: *{plan_name} Package*.\n\nPlease reply with your **M-Pesa Number** (e.g., 0712345678) to initiate secure payment via STK Push."
                    else:
                        reply_text = f"Invalid selection. Please choose a package to continue:\n\n{PRICING_MENU}"
                
                elif step == 'confirm_phone':
                    target_phone = incoming_msg.strip().replace(' ', '').replace('+', '')
                    if len(target_phone) < 10:
                        reply_text = "Please enter a valid M-Pesa number (e.g., 0712345678)."
                    else:
                        plan_id = pay_state.get('plan_id')
                        amount = pay_state.get('amount')
                        try:
                            res = initiate_stk_push(target_phone, int(amount))
                            if res and res.get('ResponseCode') == '0':
                                checkout_id = res.get("CheckoutRequestID")
                                rtdb.reference(f'pending_transactions/{checkout_id}').set({
                                    'user_id': f"wa_{user_phone}", 'amount': amount, 'plan_id': plan_id,
                                    'status': 'awaiting_payment', 'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                })
                                reply_text = f"🔄 *Initiating Secure M-Pesa Payment...*\n\nPlease enter your PIN on your phone. Your {plan_id.capitalize()} intelligence will be unlocked instantly."
                                pay_ref.delete()
                            else:
                                reply_text = "❌ M-Pesa initiation failed. Please ensure the number is correct or try again by typing 'Upgrade'."
                                pay_ref.delete()
                        except:
                            reply_text = "⚠️ Payment gateway is temporarily busy. Please try again shortly."
                            pay_ref.delete()
            
            resp = MessagingResponse(); resp.message(reply_text); return str(resp)

        # 4. ACTIVE SUBSCRIBER LOGIC
        # Greetings
        if any(greet in incoming_msg.lower() for greet in GREETINGS):
            reply_text = f"{persona_intro}Welcome back. You are currently on the *{user_tier.capitalize()}* plan.\n\n"
            reply_text += "🔍 *Prices:* Send a crop name (e.g., 'Maize')\n"
            reply_text += "🧠 *Strategy:* Send 'Farmer Advice' or 'Action Plan'"
            resp = MessagingResponse(); resp.message(reply_text); return str(resp)

        # Market Analysis & Advice
        market_ref = rtdb.reference('market_data')
        all_market = market_ref.get() or {}
        zim_items = [item for item in all_market.values() if item.get('country') == 'Zimbabwe']
        
        found_item = None
        for item in zim_items:
            commodity = item.get('commodity', '').lower()
            if commodity in incoming_msg.lower() or incoming_msg.lower() in commodity:
                found_item = item
                break

        if found_item:
            commodity_name = found_item['commodity']
            price = found_item['price']
            currency = found_item['currency']
            unit = found_item['unit']
            trend = found_item.get('trend', 'stable')
            region = found_item.get('region', 'Zimbabwe Hubs')
            date = found_item.get('updated_at', '')[:10]

            advice = ""
            if trend == 'up':
                advice = f"📈 *Strategy:* {commodity_name} prices are rising. *Action:* Sell in batches now to capitalize on high demand peaks."
            elif trend == 'down':
                advice = f"📉 *Strategy:* Prices are falling. *Action:* Hold and store. Avoid selling in a crashed market if possible."
            else:
                advice = f"⚖️ *Insight:* Market is stable. Safe window for standard trading."

            reply_text = f"{persona_intro}*Market Data: {commodity_name}*\n\n"
            reply_text += f"📍 *Market:* {region}\n"
            reply_text += f"💰 *Price:* {currency} {price:,.2f} per {unit}\n"
            reply_text += f"📊 *Trend:* {'Rising ▲' if trend == 'up' else 'Dropping ▼' if trend == 'down' else 'Stable ▬'}\n"
            reply_text += f"📅 *As of:* {date}\n\n"
            reply_text += f"{advice}"

        elif 'farmer advice' in incoming_msg.lower():
            reply_text = f"{persona_intro}🚜 *Farmer Strategic Intelligence*\n\n1️⃣ Selection: Pivot based on forecasted deficits.\n2️⃣ Timing: Historical spike detection suggests holding stock for 3 months post-harvest.\n3️⃣ Logistics: Direct-to-retail bypasses middleman costs."
        
        elif 'action plan' in incoming_msg.lower():
            reply_text = f"{persona_intro}📋 *Market Action Plan*\n\n📈 *Shortage:* Sell incrementally to ride the price curve.\n📉 *Glut:* Focus on value-add (drying/pulping) to preserve margins."
            
        else:
            reply_text = f"{persona_intro}I could not locate intelligence for '{incoming_msg}'. Please specify a commodity (e.g., 'Maize') or say 'Menu'."

    except Exception as e:
        print(f"Zim Bot Twilio Error: {e}")
        reply_text = "🇿🇼 *ZimBot Notice*\n\nI'm briefly offline for market recalibration. Please try again in 5 minutes."

    resp = MessagingResponse()
    resp.message(reply_text)
    return str(resp)

# Rule-Based Broadcast Generator
def generate_zim_market_alert(affected_crops_info):
    """
    affected_crops_info: list of dicts with {'commodity': '...', 'trend': '...', 'price': '...'}
    """
    alert_text = "⚠️ *ALERT: ZIMBOT MARKET UPDATE* ⚠️\n\n"
    alert_text += "Significant price shifts detected in the Zimbabwe agricultural network:\n\n"
    
    for crop in affected_crops_info:
        trend = crop.get('trend', 'stable')
        commodity = crop.get('commodity')
        price = crop.get('price')
        
        if trend == 'up':
            alert_text += f"🚀 *{commodity}:* Prices are RISING! Now at {price}. Sell now to maximize profit.\n"
        elif trend == 'down':
            alert_text += f"📉 *{commodity}:* Prices are DROPPING! Now at {price}. Hold your stock if possible.\n"
        else:
            alert_text += f"⚖️ *{commodity}:* Prices are STABLE at {price}.\n"
            
    alert_text += "\n*Action:* Review your trading strategy immediately based on these insights.\n\n_Stay informed with ZimBot._"
    return alert_text

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


if __name__ == "__main__":
    import os
    import socket
    port = int(os.environ.get("PORT", 5000))
    
    # Self-healing dynamic port check to prevent address collisions (WinError 10048)
    def is_port_in_use(port_to_check):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            try:
                s.bind(('0.0.0.0', port_to_check))
                return False
            except socket.error:
                return True
                
    while is_port_in_use(port):
        print(f" Port {port} is currently in use. Fallback to next address...")
        port += 1
        
    print(f"\n Farmerman Systems is LIVE!")
    print(f" Click here to open: http://127.0.0.1:{port}\n")
    
    socketio.run(app, host='0.0.0.0', port=port, debug=True)
    
    